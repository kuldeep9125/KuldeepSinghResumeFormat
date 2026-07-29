from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
REPO_R08 = REPO_ROOT / "resumes" / "R08"
LOCAL_ROOT = Path("/Users/kuldeepsingh/Downloads/RESUME/R08_Enhanced_EV_Technical_Manager_Resume")

VERSION = "R08"
FILE_BASE = "Kuldeep_Singh_Enhanced_EV_Technical_Manager_Resume_R08"

NAME = "Kuldeep Singh"
HEADLINE = "EV Technical Manager | Vehicle Controls, BMS & Electrified Powertrain Integration"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"
COMPANY = "IX Energy Pvt. Ltd., Noida, India"
ROLE = "Technical Manager - Product Development"
DATES = "Jul 2018 - Present"
EDUCATION = "B.Tech, Mechanical Engineering | VIT University, Vellore | 2014 - 2018 | First Class, 75%"

SUMMARY = (
    "EV Technical Manager with 8 years of experience delivering electric and hybrid vehicle platforms from business "
    "need and system architecture through ECU software, vehicle integration, validation, homologation, and production "
    "readiness. Combines hands-on expertise in MATLAB/Simulink/Stateflow, eVCU/BMS controls, CAN/J1939 diagnostics, "
    "high-voltage batteries, charging, and power electronics with leadership of multidisciplinary engineering teams "
    "and suppliers. Delivered ARAI/ICAT-certified commercial EV and hybrid programmes, including a validated 25% "
    "fuel-efficiency improvement."
)

LEADERSHIP_SNAPSHOT = [
    "End-to-end programme ownership: translate business requirements into architecture, SORs, development plans, supplier deliverables, validation evidence, certification outputs, and release readiness.",
    "Cross-functional leadership: lead Manufacturing, Electrical & Electronics, and Mechanical Design with 6 direct and 6 indirect team members; report project initiation, priorities, risks, and readiness to the Business Development Lead and Company Director.",
    "Technical depth with delivery focus: bridge vehicle-level performance, ECU software, HV safety, component interfaces, manufacturing constraints, cost, timing, and homologation requirements.",
]

EXPERIENCE = [
    "Lead new EV and P4 hybrid vehicle-platform development across requirements, system architecture, ECU application software, HV battery engineering, power electronics, supplier development, vehicle integration, validation, and homologation.",
    "Own project initiation inputs by identifying customer and business needs, defining the next development stage, estimating technical scope, and presenting programme readiness and key decisions for leadership approval.",
    "Develop and review BMS/eVCU/VCU firmware and application software in MATLAB/Simulink/Stateflow for drive, charge, assist, regenerative braking, torque/speed control, diagnostics, derating, precharge, active discharge, fault response, and safe-state behaviour.",
    "Architect and package 5 kWh to 300 kWh battery systems across 72V-800V LFP, NMC, LTO, and ultra-capacitor technologies, covering BMS/BMU logic, contactor strategy, precharge, charging handshake, thermal management, protection, and vehicle integration.",
    "Define SORs, technical specifications, CAN/DBC interfaces, DVP expectations, acceptance criteria, and release evidence for motors, batteries, PDU, DCDC, OBC, chargers, gearboxes, controllers, HVAC, telematics, instrumentation, and embedded electronics.",
    "Drive bench and vehicle issue resolution using CAN diagnostics, J1939/DBC review, MF4 and test-data analysis, calibration, DFMEA, root-cause analysis, supplier reviews, and structured closure tracking.",
    "Govern programme risks, design maturity, supplier readiness, validation status, change control, and certification evidence while balancing performance, range, thermal limits, safety, serviceability, cost, timing, and manufacturability.",
]

SELECTED_IMPACT = [
    "16T P4 hybrid bus: delivered an ICAT-certified platform with ultra-capacitor energy storage and a validated 25% fuel-efficiency improvement through controls integration, pilot-fleet trials, issue closure, and certification readiness.",
    "6.5T commercial EV: delivered an ARAI-homologated platform with an 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km certified range; contributed eVCU/BMS logic, DCDC/PDU/OBC integration, diagnostics, HV safety, and validation.",
    "5T commercial EV: engineered vehicle and HV-system integration for an 80 kW powertrain and 56 kWh / 310V NMC battery with a 140 km range target, coordinating interfaces, suppliers, power electronics, and vehicle validation.",
    "Certified energy-storage portfolio: supported 53 kWh/332V LFP, 11.5 kWh/332V LTO, 28 kWh NMC, and 0.5 kWh/400V ultra-capacitor systems from design and controls through integration and homologation evidence.",
]

CONTROLS_AND_VALIDATION = [
    "Independently developed and validated a hybrid-control model in MATLAB/Simulink/Stateflow, including model revisions, generated target outputs, calibration artefacts, and bench/vehicle evidence.",
    "Implemented torque and speed requests, assist/regen, SOC and charge limits, thermal derating, contactor/precharge sequencing, MCU enable, fault reset, active discharge, diagnostics, telematics status, and safe-state logic.",
    "Built CAN/J1939 and DBC-based communication among VCU, BMS, motor controller, charger, telematics, and vehicle systems; used MF4 logs and calibration evidence to tune controls and support release decisions.",
    "Developed and integrated LV/HV PDU functions, STM32F4/F1 instrument-cluster electronics, a dual-CAN telematics unit, and an ultra-capacitor monitoring/control system with CAN-based diagnostics.",
]

CORE_EXPERTISE = [
    ("Vehicle & powertrain", "EV/P4 hybrid architecture, e-drive/EDU, eVCU/VCU, BMS, HV battery, charging, DCDC, OBC, PDU, motor controller, BTMS/HVAC, vehicle integration, homologation"),
    ("Software & controls", "MATLAB, Simulink, Stateflow, Model-Based Design, Embedded C/C++, ECU firmware/application software, control logic, calibration, generated outputs, A2L/MOT artefacts"),
    ("Networks & diagnostics", "CAN, J1939, DBC, Vector CANoe, Peak CAN, BusMaster, CSS Electronics, MF4 logs, UDS awareness, DTC/SPN diagnostics, signal and interface analysis"),
    ("Systems & validation", "requirements, SOR, V-model lifecycle, interface control, DVP, MIL/SIL/HIL readiness, bench/vehicle testing, traceability, release documentation, ARAI/ICAT evidence"),
    ("Quality & leadership", "ISO 26262 awareness, IATF 16949, DFMEA, RCA, Six Sigma/DMAIC, supplier development, risk reviews, change control, team mentoring, production readiness"),
]

CERTIFICATIONS = [
    "Six Sigma Black Belt - Certified",
    "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
    "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
    "AWS IoT & Simulink - Amazon Web Services / MathWorks Certified",
]

IMPROVEMENT_NOTES = """# R08 Enhancement Notes

## Positioning

R08 positions Kuldeep Singh as an EV Technical Manager who combines business-to-engineering ownership, people leadership, hands-on vehicle-control software, electrified powertrain integration, and homologation delivery.

## Improvements from R07

- Added the confirmed reporting line to the Business Development Lead and Company Director.
- Added leadership scope across Manufacturing, Electrical & Electronics, and Mechanical Design, including 6 direct and 6 indirect team members.
- Strengthened project-initiation and business-requirement ownership.
- Consolidated repeated technical sections into a clearer two-page recruiter and hiring-manager narrative.
- Reworked bullets to lead with ownership, action, technical scope, and measurable outcomes.
- Preserved privacy-clean wording and avoided internal project codes, file names, and confidential implementation details.
- Retained ATS coverage for EV systems, ECU software, MATLAB/Simulink/Stateflow, BMS/eVCU/VCU, CAN/J1939/DBC, HV batteries, validation, ISO 26262, and ARAI/ICAT homologation.

## Recommended Use

Use R08 as the primary general resume for EV Technical Manager, EV Systems Lead, Vehicle Controls Lead, BMS/eVCU Lead, Electrified Powertrain Integration, and automotive product-development roles.
"""


NAVY = "0B2545"
BLUE = "1F4D78"
GRAY = "5B6573"
LIGHT_BLUE = "DCE8F2"
RULE = "C7D2DC"


def set_font(run, name="Arial", size=9.1, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=2.2, line=1.06):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def set_keep(paragraph, keep_next=False, keep_together=True):
    fmt = paragraph.paragraph_format
    fmt.keep_with_next = keep_next
    fmt.keep_together = keep_together


def add_bottom_rule(paragraph, color=RULE, size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_numbering(doc):
    numbering = doc.part.numbering_part.element
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), "42")
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "310")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "310")
    ind.set(qn("w:hanging"), "180")
    p_pr.append(tabs)
    p_pr.append(ind)
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    r_pr.append(color)
    for element in [start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr]:
        lvl.append(element)
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), "42")
    abstract_id = OxmlElement("w:abstractNumId")
    abstract_id.set(qn("w:val"), "42")
    num.append(abstract_id)
    numbering.append(num)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.46)
    section.bottom_margin = Inches(0.44)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.1)
    normal.paragraph_format.space_after = Pt(2.2)
    normal.paragraph_format.line_spacing = 1.06

    if "Resume Bullet" not in [s.name for s in doc.styles]:
        bullet_style = doc.styles.add_style("Resume Bullet", WD_STYLE_TYPE.PARAGRAPH)
    else:
        bullet_style = doc.styles["Resume Bullet"]
    bullet_style.base_style = normal
    bullet_style.font.name = "Arial"
    bullet_style.font.size = Pt(8.95)
    bullet_style.paragraph_format.space_after = Pt(1.35)
    bullet_style.paragraph_format.line_spacing = 1.05

    configure_numbering(doc)


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, after=0.5, line=1.0)
    set_keep(title, keep_next=True)
    set_font(title.add_run(NAME.upper()), size=18.2, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=0.7, line=1.0)
    set_keep(subtitle, keep_next=True)
    set_font(subtitle.add_run(HEADLINE), size=9.7, bold=True, color=BLUE)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=2.9, line=1.0)
    set_font(contact.add_run(CONTACT), size=8.25, color=GRAY)


def add_heading(doc, text, before=5.2):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=2.2, line=1.0)
    set_keep(p, keep_next=True)
    set_font(p.add_run(text), size=10.4, bold=True, color=BLUE)
    add_bottom_rule(p)
    return p


def add_body(doc, text, size=9.05, after=1.7):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.06)
    set_font(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=8.9, after=1.25):
    p = doc.add_paragraph(style="Resume Bullet")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "42")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    set_spacing(p, after=after, line=1.05)
    set_font(p.add_run(text), size=size)
    return p


def add_role_header(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0.4, after=0.15, line=1.0)
    set_keep(p, keep_next=True)
    set_font(p.add_run(ROLE), size=9.35, bold=True, color=NAVY)
    p2 = doc.add_paragraph()
    set_spacing(p2, after=1.2, line=1.0)
    set_keep(p2, keep_next=True)
    set_font(p2.add_run(f"{COMPANY} | {DATES}"), size=8.35, italic=True, color=GRAY)


def add_expertise_line(doc, label, detail):
    p = doc.add_paragraph()
    set_spacing(p, after=1.45, line=1.05)
    set_font(p.add_run(f"{label}: "), size=8.85, bold=True, color=NAVY)
    set_font(p.add_run(detail), size=8.85)


def set_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_font(run, size=7.8, color=GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def build_resume(out_dir: Path):
    doc = Document()
    configure_doc(doc)
    set_page_number(doc.sections[0].footer.paragraphs[0])
    add_header(doc)

    add_heading(doc, "EXECUTIVE PROFILE", before=1.2)
    add_body(doc, SUMMARY, size=9.05, after=1.4)

    add_heading(doc, "LEADERSHIP & DELIVERY SNAPSHOT")
    for item in LEADERSHIP_SNAPSHOT:
        add_bullet(doc, item, size=8.95)

    add_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_header(doc)
    for item in EXPERIENCE:
        add_bullet(doc, item, size=8.84, after=1.12)

    add_heading(doc, "SELECTED PROGRAMME IMPACT")
    for item in SELECTED_IMPACT:
        add_bullet(doc, item, size=8.86, after=1.18)

    doc.add_page_break()
    add_heading(doc, "MODEL-BASED CONTROLS, FIRMWARE & VALIDATION", before=0)
    for item in CONTROLS_AND_VALIDATION:
        add_bullet(doc, item, size=8.88, after=1.2)

    add_heading(doc, "CORE TECHNICAL EXPERTISE")
    for label, detail in CORE_EXPERTISE:
        add_expertise_line(doc, label, detail)

    add_heading(doc, "CERTIFICATIONS")
    for item in CERTIFICATIONS:
        add_bullet(doc, item, size=8.88, after=0.75)

    add_heading(doc, "EDUCATION")
    add_body(doc, EDUCATION, size=8.9, after=0)

    docx_path = out_dir / f"{FILE_BASE}.docx"
    doc.save(docx_path)
    return docx_path


def build_text(out_dir: Path):
    sections = [
        NAME.upper(), HEADLINE, CONTACT, "", "EXECUTIVE PROFILE", SUMMARY, "",
        "LEADERSHIP & DELIVERY SNAPSHOT", *[f"- {x}" for x in LEADERSHIP_SNAPSHOT], "",
        "PROFESSIONAL EXPERIENCE", ROLE, f"{COMPANY} | {DATES}", *[f"- {x}" for x in EXPERIENCE], "",
        "SELECTED PROGRAMME IMPACT", *[f"- {x}" for x in SELECTED_IMPACT], "",
        "MODEL-BASED CONTROLS, FIRMWARE & VALIDATION", *[f"- {x}" for x in CONTROLS_AND_VALIDATION], "",
        "CORE TECHNICAL EXPERTISE", *[f"{label}: {detail}" for label, detail in CORE_EXPERTISE], "",
        "CERTIFICATIONS", *[f"- {x}" for x in CERTIFICATIONS], "", "EDUCATION", EDUCATION, "",
    ]
    txt_path = out_dir / f"{FILE_BASE}.txt"
    txt_path.write_text("\n".join(sections), encoding="utf-8")
    return txt_path


def build_notes(out_dir: Path):
    readme = f"""# R08 Enhanced EV Technical Manager Resume

Primary two-page resume for EV Technical Manager, Vehicle Controls Lead, BMS/eVCU Lead, EV Systems Lead, and Electrified Powertrain Integration roles.

## Files

- `{FILE_BASE}.docx`: editable Word resume
- `{FILE_BASE}.pdf`: application-ready PDF resume
- `{FILE_BASE}.txt`: ATS/plain-text resume
- `R08_Enhancement_Notes.md`: positioning and improvement summary

Version: `{VERSION}`
"""
    (out_dir / "README.md").write_text(readme, encoding="utf-8")
    (out_dir / "R08_Enhancement_Notes.md").write_text(IMPROVEMENT_NOTES, encoding="utf-8")


def sync_local(out_dir: Path):
    if LOCAL_ROOT.exists():
        shutil.rmtree(LOCAL_ROOT)
    LOCAL_ROOT.mkdir(parents=True, exist_ok=True)
    for file in out_dir.iterdir():
        if file.is_file():
            shutil.copy2(file, LOCAL_ROOT / file.name)


def main():
    if REPO_R08.exists():
        shutil.rmtree(REPO_R08)
    REPO_R08.mkdir(parents=True, exist_ok=True)
    build_resume(REPO_R08)
    build_text(REPO_R08)
    build_notes(REPO_R08)
    sync_local(REPO_R08)
    print(REPO_R08)
    print(LOCAL_ROOT)


if __name__ == "__main__":
    main()
