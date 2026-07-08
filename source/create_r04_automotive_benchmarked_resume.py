from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kuldeepsingh/Documents/Codex/2026-07-08/en")
OUTPUT_ROOT = ROOT / "outputs" / "R04_Automotive_Benchmarked_Lead_Resume"
REPO_ROOT = ROOT / "work" / "KuldeepSinghResumeFormat"
REPO_R04 = REPO_ROOT / "resumes" / "R04"
LOCAL_ROOT = Path("/Users/kuldeepsingh/Downloads/RESUME/R04_Automotive_Benchmarked_Lead_Resume")

VERSION = "R04"
FILE_BASE = "Kuldeep_Singh_Automotive_EV_Software_Technical_Lead_Resume_R04"

NAME = "Kuldeep Singh"
HEADLINE = "Automotive EV Software & Technical Lead - eVCU/BMS, MBD, HV Battery, Vehicle Integration"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"
COMPANY = "IX Energy Pvt. Ltd., Noida, India"
ROLE = "Technical Manager - Product Development"
DATES = "Jul 2018 - Present"
EDUCATION = "B.Tech - Mechanical Engineering, VIT University, Vellore, Tamil Nadu | 2014 - 2018 | First Class, 75%"

SUMMARY = (
    "Automotive EV software and technical lead with 8 years taking commercial EV, hybrid, HV battery, "
    "embedded electronics, and vehicle integration programmes from requirements and architecture through "
    "validation, release evidence, and ICAT/ARAI homologation. Strong fit for lead roles that need a hands-on "
    "manager across MATLAB/Simulink/Stateflow model-based development, eVCU/BMS application software, HV "
    "battery systems, DCDC/OBC/PDU integration, CAN diagnostics, supplier issue closure, V-model traceability, "
    "functional-safety discipline, and cross-functional product delivery. Brings practical lead-role proof: "
    "certified commercial EV platforms, 5 kWh to 300 kWh battery systems, AC/DC charging logic, "
    "embedded vehicle electronics, and quantified hybrid fuel-efficiency validation."
)

TARGET_ROLES = [
    "Automotive EV Software Lead / eVCU-BMS Lead / Vehicle Controls Lead",
    "Technical Manager - EV Powertrain, Battery Systems, Embedded Electronics",
    "Systems Integration Lead - HV Battery, Power Electronics, Charging, Vehicle Validation",
    "MBD / Controls / Validation Lead for OEM, Tier-1, EV startup, and engineering-services roles",
]

SCREENING_PROOF = [
    "Vehicle software ownership: eVCU/BMS operating states, drive readiness, charging mode, DCDC/OBC/PDU coordination, diagnostics, fault response, power sequencing, and safe shutdown behaviour.",
    "Systems leadership: architecture, SOR, supplier interfaces, component readiness, software/hardware integration, test planning, validation evidence, and homologation handoff.",
    "Validation depth: CAN diagnostics, model checks, range estimation, vehicle-level validation, calibration behaviour review, DFMEA/RCA, issue closure, certification evidence, and release-quality documentation.",
    "Automotive process fit: V-model delivery, requirements traceability, ISO 26262 awareness, IATF 16949 quality mindset, Six Sigma Black Belt methods, AUTOSAR basics, and SDV readiness.",
]

EXPERIENCE = [
    "Lead EV and P4 hybrid product development across powertrain architecture, ECU application software, HV battery engineering, supplier development, vehicle integration, validation, and homologation readiness.",
    "Develop BMS and eVCU application software in MATLAB/Simulink/Stateflow, translating functional requirements and safety constraints into model-based logic for drive, charge, fault, diagnostic, and safe-state behaviour.",
    "Define SORs, component specifications, interface requirements, test evidence, and supplier deliverables for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, EPS, BCS, TCS, HVAC, telematics, instrumentation, and embedded electronics.",
    "Architect certified commercial EV systems from requirement capture through ICAT/ARAI certification outputs, linking hardware, software, supplier readiness, validation, and vehicle-level acceptance evidence.",
    "Design and package 5 kWh to 300 kWh HV battery systems across 72V-800V LFP, NMC, LTO, and ultra-capacitor chemistries, including BMU logic, contactor/precharge interfaces, BTMS/HVAC validation, charging handshake, and pack integration.",
    "Build EV/hybrid range estimation and control models; support MIL/SIL-style model review, bench/vehicle testing, CAN diagnostics, fault investigation, calibration behaviour review, and validation-cycle reduction.",
    "Manage internal engineering teams and external suppliers across mechanical, electrical, embedded software, battery, power electronics, quality, and vehicle integration workstreams; communicate risk, readiness, and issue closure to leadership.",
]

PROGRAMMES = [
    "6.5T LCV/LPT commercial EV: ARAI-homologated platform with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed eVCU/BMS logic, vehicle control, DCDC/PDU/OBC integration, charging interfaces, diagnostics, safety logic, and validation support.",
    "5T commercial EV: supported 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, HV battery integration, power electronics interfaces, supplier coordination, and vehicle validation.",
    "16T P4 hybrid bus: ICAT-certified hybrid platform with super-capacitor energy storage and validated 25% fuel-efficiency improvement; supported system integration, pilot fleet trials, validation evidence, and certification readiness.",
    "Multi-platform EV conversions: supported sedan and light commercial EV conversions across HV battery, charger, DCDC, controller, vehicle controls, diagnostics, and validation integration.",
]

PORTFOLIO = [
    "Battery and charging: 53 kWh/332V LFP electric truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC pack, 0.5 kWh/400V ultra-capacitor pack, AC/DC charging controller logic, charger/BMS/VCU coordination, HV-LV handshake, and fault response.",
    "Embedded electronics: LV/HV PDU, STM32F4/F1 instrument cluster, Quectel dual-CAN telematics unit, ultra-capacitor cell monitoring/control system, battery-related software interfaces, and CAN-based diagnostics.",
    "Tools and methods: MATLAB, Simulink, Stateflow, Model-Based Design, Embedded C/C++, STM32, Microchip PIC, NXP, Infineon exposure, DBC, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, DFMEA, RCA, DMAIC, SPC, control charts, SOR, V-model, ICAT, ARAI.",
]

TECHNICAL_TOOLKIT = [
    "Automotive software: eVCU, VCU, BMS, vehicle controls, operating-state logic, diagnostics, fault handling, calibration, validation, release evidence, software/hardware integration.",
    "EV systems: HV battery, LFP, NMC, LTO, ultra-capacitor, DCDC, OBC, PDU, charger, contactor, precharge, power conversion, BTMS/HVAC, e-drive/EDU, commercial EV, hybrid powertrain.",
    "Network and validation: CAN, DBC, Vector CANoe, Peak CAN, Bus Master, UDS-aware diagnostics, MIL/SIL/HIL readiness, test planning, instrumentation, validation evidence, homologation.",
    "Leadership and process: technical manager, supplier development, cross-functional leadership, requirements traceability, V-model, DFMEA, RCA, ISO 26262, IATF 16949, Six Sigma, AUTOSAR basics, SDV awareness.",
]

CERTIFICATIONS = [
    "Six Sigma Black Belt - Certified",
    "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
    "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
    "AWS IoT & SIMULINK - Amazon Web Services / MathWorks Certified",
]

BENCHMARK_NOTES = [
    "# R04 Automotive Resume Benchmark Notes",
    "",
    "## Verdict",
    "",
    "R03 was strong, but not fully optimized for one-shot automotive screening because it underplayed several phrases repeatedly visible in current EV/automotive software lead roles: requirements traceability, V-model evidence, MIL/SIL/HIL readiness, diagnostics, AUTOSAR/SDV awareness, software/hardware integration, supplier issue closure, and release-quality validation evidence.",
    "",
    "R04 improves the resume without inventing unsupported experience. It uses stronger lead-role positioning, brings proof earlier, and adds market language where Kuldeep's provided background supports it.",
    "",
    "## Benchmark Signals Used",
    "",
    "- Tesla-style BMS, charging, power conversion, and vehicle software roles emphasize firmware/control logic, validation, diagnostics, and integration across physical systems.",
    "- Daimler/Mercedes-style commercial vehicle software roles emphasize MATLAB, Vector tools, embedded software, AUTOSAR, CAN/LIN/Ethernet, ISO 26262, SDV direction, and global engineering coordination.",
    "- Tata Technologies embedded engineering messaging emphasizes V-cycle delivery, MBD, MIL/SIL/HIL, AUTOSAR, SDV, and automotive engineering services.",
    "- JLR EDU/powertrain validation roles emphasize component/subsystem/system testing, issue resolution, MATLAB/data workflows, DVP/test evidence, and cross-functional validation.",
    "- AVL e-mobility and simulation roles emphasize electrification, simulation, test/validation systems, charging, battery/e-drive/inverter validation, and virtual development.",
    "- Bosch/Continental mobility roles emphasize embedded software, hardware/software integration, system engineering, automotive applications, safety/process discipline, and mobility product delivery.",
    "",
    "## Changes Made In R04",
    "",
    "- Changed headline to a broader automotive lead positioning instead of generic lead positioning.",
    "- Added a `CORE AUTOMOTIVE LEADERSHIP PROOF` section on page 1 so match evidence appears faster.",
    "- Rewrote current-role bullets around requirements-to-release flow, V-model traceability, validation evidence, supplier deliverables, and systems ownership.",
    "- Added MIL/SIL-style model review, HIL readiness, UDS-aware diagnostics, AUTOSAR basics, and SDV awareness carefully as readiness/awareness, not false production ownership.",
    "- Preserved quantified proof: 6.5T EV, 5T EV, 16T hybrid bus, 5 kWh to 300 kWh HV battery systems, and 25% fuel-efficiency validation.",
    "",
    "## Research Sources",
    "",
    "- https://www.tesla.com/careers/search",
    "- https://www.tesla.com/careers/search/job/embedded-firmware-engineer-battery-management-system-254793",
    "- https://www.tesla.com/careers/search/job/software-integration-engineer-power-conversion-charging-255576",
    "- https://dtici.daimlertruck.com/career/",
    "- https://jobsearch.daimlertruck.com/",
    "- https://www.tatatechnologies.com/us/solutions/",
    "- https://www.jaguarlandrovercareers.com/",
    "- https://www.avl.com/en/career/fields-interest/e-mobility",
    "- https://www.bosch.com/careers/",
    "- https://www.continental.com/en/career/",
    "",
]


def clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def set_font(run, name="Arial", size=9.0, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=2.2, line=1.02):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_rule(paragraph, color="D9D9D9", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    for style_name, size in [("Normal", 9.0), ("List Bullet", 8.72)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(2.2)
        style.paragraph_format.line_spacing = 1.02


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=0, after=1, line=1.0)
    set_font(title.add_run(NAME.upper()), size=17.2, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=1, line=1.0)
    set_font(subtitle.add_run(HEADLINE), size=9.7, bold=True, color="1F4D78")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=4, line=1.0)
    set_font(contact.add_run(CONTACT), size=8.3, color="555555")


def add_continuation_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=4, line=1.0)
    set_font(p.add_run("KULDEEP SINGH | Automotive EV Software & Technical Lead | Page 2"), size=8.3, bold=True, color="555555")


def add_heading(doc, text, before=6):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=2.5, line=1.0)
    set_font(p.add_run(text), size=10.3, bold=True, color="1F4D78")
    add_bottom_rule(p)


def add_body(doc, text, size=8.95, after=2.4):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.02)
    set_font(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=8.62, after=1.0):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=after, line=1.0)
    p.paragraph_format.left_indent = Inches(0.21)
    p.paragraph_format.first_line_indent = Inches(-0.13)
    set_font(p.add_run(text), size=size)
    return p


def add_role_header(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=1, after=0, line=1.0)
    set_font(p.add_run(ROLE), size=9.55, bold=True)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=1.6, line=1.0)
    set_font(p2.add_run(f"{COMPANY} | {DATES}"), size=8.55, color="555555")


def build_resume(out_dir: Path):
    doc = Document()
    configure_doc(doc)
    add_header(doc)

    add_heading(doc, "EXECUTIVE SUMMARY", before=2)
    add_body(doc, SUMMARY)

    add_heading(doc, "TARGET LEAD ROLES")
    for item in TARGET_ROLES:
        add_bullet(doc, item, size=8.68)

    add_heading(doc, "CORE AUTOMOTIVE LEADERSHIP PROOF")
    for item in SCREENING_PROOF:
        add_bullet(doc, item, size=8.62, after=0.8)

    add_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_header(doc)
    for item in EXPERIENCE:
        add_bullet(doc, item, size=8.56, after=0.8)

    doc.add_page_break()

    add_continuation_header(doc)
    add_heading(doc, "SELECTED PROGRAMMES & QUANTIFIED PROOF", before=0)
    for item in PROGRAMMES:
        add_bullet(doc, item, size=8.62, after=0.9)

    add_heading(doc, "BATTERY, CHARGING, EMBEDDED & VALIDATION PORTFOLIO")
    for item in PORTFOLIO:
        add_bullet(doc, item, size=8.58, after=0.9)

    add_heading(doc, "TECHNICAL TOOLKIT & DELIVERY METHODS")
    for item in TECHNICAL_TOOLKIT:
        add_bullet(doc, item, size=8.48, after=0.65)

    add_heading(doc, "CERTIFICATIONS")
    for item in CERTIFICATIONS:
        add_bullet(doc, item, size=8.62, after=0.6)

    add_heading(doc, "EDUCATION")
    add_body(doc, EDUCATION, size=8.75, after=0)

    docx_path = out_dir / f"{FILE_BASE}.docx"
    txt_path = out_dir / f"{FILE_BASE}.txt"
    doc.save(docx_path)

    txt_lines = [
        NAME.upper(),
        HEADLINE,
        CONTACT,
        "",
        "EXECUTIVE SUMMARY",
        SUMMARY,
        "",
        "TARGET LEAD ROLES",
        *[f"- {x}" for x in TARGET_ROLES],
        "",
        "CORE AUTOMOTIVE LEADERSHIP PROOF",
        *[f"- {x}" for x in SCREENING_PROOF],
        "",
        "PROFESSIONAL EXPERIENCE",
        ROLE,
        f"{COMPANY} | {DATES}",
        *[f"- {x}" for x in EXPERIENCE],
        "",
        "SELECTED PROGRAMMES & QUANTIFIED PROOF",
        *[f"- {x}" for x in PROGRAMMES],
        "",
        "BATTERY, CHARGING, EMBEDDED & VALIDATION PORTFOLIO",
        *[f"- {x}" for x in PORTFOLIO],
        "",
        "TECHNICAL TOOLKIT & DELIVERY METHODS",
        *[f"- {x}" for x in TECHNICAL_TOOLKIT],
        "",
        "CERTIFICATIONS",
        *[f"- {x}" for x in CERTIFICATIONS],
        "",
        "EDUCATION",
        EDUCATION,
        "",
    ]
    txt_path.write_text("\n".join(txt_lines), encoding="utf-8")
    return docx_path, txt_path


def build_readme(out_dir: Path):
    readme = [
        "# R04 Automotive-Benchmarked Lead Resume",
        "",
        "Automotive-industry optimized resume for broad EV software, technical manager, systems integration, battery, embedded, validation, and supplier-leadership roles.",
        "",
        "## Files",
        f"- `{FILE_BASE}.docx`: editable Word resume",
        f"- `{FILE_BASE}.pdf`: rendered PDF resume",
        f"- `{FILE_BASE}.txt`: ATS/plain-text resume",
        "- `R04_Benchmark_Notes.md`: research-backed comparison and improvement notes",
        "",
        "## Why R04 Exists",
        "- R03 was strong but generic.",
        "- R04 moves market screening proof to page 1 and aligns language with current automotive EV software lead requirements.",
        "- ATS terms are woven into normal resume sections rather than placed in a standalone keyword section.",
        "",
        f"Version: `{VERSION}`",
        "",
    ]
    (out_dir / "README.md").write_text("\n".join(readme), encoding="utf-8")
    (out_dir / "R04_Benchmark_Notes.md").write_text("\n".join(BENCHMARK_NOTES), encoding="utf-8")


def main():
    for path in [OUTPUT_ROOT, REPO_R04, LOCAL_ROOT]:
        clean_dir(path)

    docx_path, txt_path = build_resume(OUTPUT_ROOT)
    build_readme(OUTPUT_ROOT)

    for file in OUTPUT_ROOT.iterdir():
        if file.is_file():
            shutil.copy2(file, REPO_R04 / file.name)
            shutil.copy2(file, LOCAL_ROOT / file.name)

    print(OUTPUT_ROOT)
    print(docx_path)
    print(txt_path)


if __name__ == "__main__":
    main()
