from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kuldeepsingh/Documents/Codex/2026-07-08/en")
OUTPUT_ROOT = ROOT / "outputs" / "R03_Generic_Lead_Resume"
REPO_ROOT = ROOT / "work" / "KuldeepSinghResumeFormat"
REPO_R03 = REPO_ROOT / "resumes" / "R03"
LOCAL_ROOT = Path("/Users/kuldeepsingh/Downloads/RESUME/R03_Generic_Lead_Resume")

VERSION = "R03"
FILE_BASE = "Kuldeep_Singh_EV_Software_Technical_Manager_Lead_Resume_R03"

NAME = "Kuldeep Singh"
HEADLINE = "EV Software Lead | Technical Manager - eVCU/BMS, MBD, HV Battery & Vehicle Integration"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"
COMPANY = "IX Energy Pvt. Ltd., Noida, India"
ROLE = "Technical Manager - Product Development"
DATES = "Jul 2018 - Present"
EDUCATION = "B.Tech - Mechanical Engineering, VIT University, Vellore, Tamil Nadu | 2014 - 2018 | First Class, 75%"

# Design preset: compact_reference_guide, with a resume-density override.
# Exact override tokens: Letter page, 0.50 in margins, Arial, navy heading
# system, real Word paragraph styles, and List Bullet paragraphs with explicit
# indents/spacing. The page break after current role creates the requested
# two-page form while keeping ATS text linear.

SUMMARY = (
    "EV software lead and technical manager with 8 years delivering commercial EV, hybrid, "
    "battery, embedded electronics, and vehicle integration programmes from concept through "
    "validation and ICAT/ARAI homologation. Strong blend of hands-on MATLAB/Simulink/Stateflow "
    "model-based software, eVCU/BMS logic, HV battery systems, DCDC/OBC/PDU integration, CAN "
    "diagnostics, supplier coordination, and cross-functional delivery leadership. Known for "
    "turning requirements, SORs, safety constraints, and supplier interfaces into validated "
    "vehicle-level solutions for lead roles across automotive, EV software, embedded systems, "
    "powertrain integration, and technical programme execution."
)

ROLE_FIT = [
    "EV Software Lead / eVCU-BMS Lead / Vehicle Controls Lead",
    "Technical Manager - EV Powertrain, Battery Systems, Embedded Electronics",
    "Automotive Embedded Systems, Model-Based Development, and System Integration Lead",
    "Validation, Homologation, Supplier Development, and Cross-Functional Delivery Lead",
]

LEADERSHIP_STRENGTHS = [
    "Lead complete EV and P4 hybrid architecture across ECU software, HV battery, power electronics, vehicle controls, validation, supplier development, and certification readiness.",
    "Translate product requirements into SORs, functional requirements, safety requirements, model-based control logic, interface definitions, test plans, and release-ready engineering evidence.",
    "Manage internal engineering teams and external suppliers across battery, motor, gearbox, DCDC, OBC, PDU, charger, controller, EPS, BCS, TCS, HVAC, sensors, telematics, and embedded electronics.",
    "Apply DFMEA, RCA, DMAIC, SPC, control charts, Six Sigma, ISO 26262 awareness, and IATF 16949 process discipline to improve robustness, supplier quality, and validation efficiency.",
]

EXPERIENCE = [
    "Lead EV and P4 hybrid product development covering powertrain architecture, ECU application software, battery engineering, supplier development, vehicle integration, validation, and homologation readiness for commercial vehicle electrification programmes.",
    "Develop BMS and eVCU application software using MATLAB/Simulink/Stateflow for vehicle operating states, drive readiness, charging mode, DCDC/OBC/PDU coordination, diagnostics, fault response, power sequencing, and safe shutdown behaviour.",
    "Architect commercial EV systems from SOR definition through ICAT/ARAI certification outputs, connecting product requirements to hardware, software, testing, supplier deliverables, and vehicle-level evidence.",
    "Design and package HV battery systems from 5 kWh to 300 kWh across 72V-800V LFP, NMC, LTO, and ultra-capacitor chemistries, including BMU development, BTMS/HVAC validation, and pack integration.",
    "Define component and system specifications for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, EPS, BCS, TCS, HVAC, telematics, instrumentation, and embedded vehicle electronics.",
    "Build EV/hybrid range estimation models and support vehicle-level validation using CAN tools, diagnostics, fault investigation, calibration behaviour review, supplier issue closure, and homologation evidence preparation.",
    "Coordinate with senior OEM/Tier-1 leadership, suppliers, and internal teams to maintain component readiness, release-quality software/hardware interfaces, validation closure, and programme delivery discipline.",
]

DELIVERY_SCOPE = [
    "Own lead-level delivery across requirement capture, architecture, supplier SOR, software logic, component readiness, validation evidence, and certification handoff.",
    "Bridge mechanical, electrical, embedded software, battery, power electronics, quality, supplier, and vehicle integration teams with clear risk, readiness, and issue-closure communication.",
    "Balance hands-on engineering depth with technical management: model reviews, interface decisions, supplier follow-up, validation prioritisation, and release-quality documentation.",
]

PROGRAMMES = [
    "LCV/LPT commercial EV conversion: ARAI-homologated 6.5T EV with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed eVCU/BMS logic, vehicle control, DCDC/PDU/OBC integration, charging interfaces, diagnostics, safety logic, and validation support.",
    "5T commercial EV conversion: supported 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, power electronics interfaces, HV battery integration, supplier coordination, and vehicle validation.",
    "16T P4 hybrid bus: ICAT-certified hybrid platform using super-capacitor energy storage with validated 25% fuel-efficiency improvement; supported vehicle integration, pilot fleet trials, validation evidence, and certification readiness.",
    "Sedan and light commercial EV conversion programmes: supported Ambassador and Mahindra Supro EV conversion work across HV battery, charger, DCDC, controller, vehicle control, diagnostics, and validation integration.",
]

BATTERY_SYSTEMS = [
    "Delivered certified HV battery packs including 53 kWh/332V LFP electric truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC pack, and 0.5 kWh/400V ultra-capacitor hybrid bus pack.",
    "Led battery architecture across electrical, mechanical, thermal, software, and safety interfaces, including contactors, precharge, BMU logic, BTMS/HVAC coordination, charger handshake, and pack-level validation.",
    "Integrated AC/DC charging controller logic, DCDC converter, OBC, PDU, HV-LV handshake, charger/BMS/VCU coordination, fault response, and vehicle safety readiness.",
]

EMBEDDED_SYSTEMS = [
    "Developed in-house LV/HV PDU, STM32F4/F1 instrument cluster, Quectel dual-CAN telematics unit, ultra-capacitor cell monitoring/control system, and battery-related software interfaces.",
    "Worked with Embedded C/C++, STM32, Microchip PIC, NXP, Infineon exposure, DBC files, CSS Electronics tools, Vector CANoe, Peak CAN, Bus Master, diagnostics, and CAN-based vehicle communication.",
    "Built MATLAB/Simulink range estimation and control models for component-level and vehicle-level EV/hybrid scenarios, supporting virtual design decisions and validation planning.",
]

TOOLKIT_LINES = [
    "EV Software & Controls: MATLAB, Simulink, Stateflow, Model-Based Design, eVCU, VCU, BMS, operating-state logic, diagnostics, fault handling, validation, calibration.",
    "Powertrain & Battery: HV battery, LFP, NMC, LTO, ultra-capacitor, DCDC, OBC, PDU, charger, motor, gearbox, EDU/e-drive, power conversion, BTMS/HVAC, AC/DC charging.",
    "Embedded & Network: Embedded C/C++, STM32, Microchip PIC, NXP, Infineon, CAN, DBC, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, telematics, instrument cluster.",
    "Delivery & Quality: SOR, DFMEA, RCA, DMAIC, SPC, control charts, supplier development, V-model, ISO 26262, IATF 16949, ICAT, ARAI, homologation, Six Sigma.",
]

CERTIFICATIONS = [
    "Six Sigma Black Belt - Certified",
    "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
    "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
    "AWS IoT & SIMULINK - Amazon Web Services / MathWorks Certified",
]


def clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def set_font(run, name="Arial", size=9.2, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=2.4, line=1.03):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_rule(paragraph, color="D9D9D9", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.50)
    section.bottom_margin = Inches(0.50)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)
    section.header_distance = Inches(0.30)
    section.footer_distance = Inches(0.30)

    for style_name, size in [("Normal", 9.15), ("List Bullet", 8.85)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(2.4)
        style.paragraph_format.line_spacing = 1.03


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=0, after=1, line=1.0)
    set_font(title.add_run(NAME.upper()), size=17.5, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=1, line=1.0)
    set_font(subtitle.add_run(HEADLINE), size=9.9, bold=True, color="1F4D78")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=5, line=1.0)
    set_font(contact.add_run(CONTACT), size=8.35, color="555555")


def add_continuation_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=5, line=1.0)
    set_font(
        p.add_run("KULDEEP SINGH | EV Software Lead & Technical Manager | Page 2"),
        size=8.4,
        bold=True,
        color="555555",
    )


def add_heading(doc, text, before=7):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=3, line=1.0)
    set_font(p.add_run(text), size=10.6, bold=True, color="1F4D78")
    add_bottom_rule(p)
    return p


def add_body(doc, text, size=9.05, after=2.6):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.03)
    set_font(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=8.82, after=1.4):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=after, line=1.02)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    set_font(p.add_run(text), size=size)
    return p


def add_role_header(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=2, after=0, line=1.0)
    set_font(p.add_run(ROLE), size=9.8, bold=True)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=2, line=1.0)
    set_font(p2.add_run(f"{COMPANY} | {DATES}"), size=8.7, color="555555")


def add_page_break(doc):
    doc.add_page_break()


def build_resume(out_dir: Path):
    doc = Document()
    configure_doc(doc)

    add_header(doc)

    add_heading(doc, "EXECUTIVE SUMMARY", before=3)
    add_body(doc, SUMMARY)

    add_heading(doc, "LEAD ROLE FIT")
    for item in ROLE_FIT:
        add_bullet(doc, item, size=8.85)

    add_heading(doc, "CORE LEADERSHIP STRENGTHS")
    for item in LEADERSHIP_STRENGTHS:
        add_bullet(doc, item)

    add_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_header(doc)
    for item in EXPERIENCE:
        add_bullet(doc, item, size=8.78, after=1.1)

    add_heading(doc, "LEADERSHIP SCOPE & DELIVERY ENVIRONMENT")
    for item in DELIVERY_SCOPE:
        add_bullet(doc, item, size=8.78, after=1.0)

    add_page_break(doc)

    add_continuation_header(doc)
    add_heading(doc, "SELECTED PROGRAMMES & QUANTIFIED PROOF", before=0)
    for item in PROGRAMMES:
        add_bullet(doc, item, size=8.78, after=1.2)

    add_heading(doc, "HV BATTERY, CHARGING & EV SYSTEMS PORTFOLIO")
    for item in BATTERY_SYSTEMS:
        add_bullet(doc, item, size=8.78, after=1.2)

    add_heading(doc, "EMBEDDED SOFTWARE & VEHICLE ELECTRONICS")
    for item in EMBEDDED_SYSTEMS:
        add_bullet(doc, item, size=8.78, after=1.2)

    add_heading(doc, "TECHNICAL TOOLKIT")
    for item in TOOLKIT_LINES:
        add_bullet(doc, item, size=8.72, after=1.0)

    add_heading(doc, "CERTIFICATIONS")
    for item in CERTIFICATIONS:
        add_bullet(doc, item, size=8.78, after=0.8)

    add_heading(doc, "EDUCATION")
    add_body(doc, EDUCATION, size=8.9, after=0)

    docx_path = out_dir / f"{FILE_BASE}.docx"
    txt_path = out_dir / f"{FILE_BASE}.txt"
    doc.save(docx_path)

    txt_lines = [
        NAME.upper(),
        HEADLINE,
        CONTACT,
        "",
        "EXECUTIVE SUMMARY",
        SUMMARY,
        "",
        "LEAD ROLE FIT",
        *[f"- {x}" for x in ROLE_FIT],
        "",
        "CORE LEADERSHIP STRENGTHS",
        *[f"- {x}" for x in LEADERSHIP_STRENGTHS],
        "",
        "PROFESSIONAL EXPERIENCE",
        ROLE,
        f"{COMPANY} | {DATES}",
        *[f"- {x}" for x in EXPERIENCE],
        "",
        "LEADERSHIP SCOPE & DELIVERY ENVIRONMENT",
        *[f"- {x}" for x in DELIVERY_SCOPE],
        "",
        "SELECTED PROGRAMMES & QUANTIFIED PROOF",
        *[f"- {x}" for x in PROGRAMMES],
        "",
        "HV BATTERY, CHARGING & EV SYSTEMS PORTFOLIO",
        *[f"- {x}" for x in BATTERY_SYSTEMS],
        "",
        "EMBEDDED SOFTWARE & VEHICLE ELECTRONICS",
        *[f"- {x}" for x in EMBEDDED_SYSTEMS],
        "",
        "TECHNICAL TOOLKIT",
        *[f"- {x}" for x in TOOLKIT_LINES],
        "",
        "CERTIFICATIONS",
        *[f"- {x}" for x in CERTIFICATIONS],
        "",
        "EDUCATION",
        EDUCATION,
        "",
    ]
    txt_path.write_text("\n".join(txt_lines), encoding="utf-8")
    return docx_path, txt_path


def build_readme(out_dir: Path):
    readme = [
        "# R03 Generic 2-Page Lead Resume",
        "",
        "Generic, company-neutral resume for EV Software Lead and Technical Manager applications across automotive, EV, embedded software, BMS/eVCU, systems integration, supplier leadership, and product development roles.",
        "",
        "## Files",
        f"- `{FILE_BASE}.docx`: editable Word resume",
        f"- `{FILE_BASE}.pdf`: rendered PDF resume",
        f"- `{FILE_BASE}.txt`: ATS/plain-text resume",
        "",
        "## Positioning",
        "- EV Software Lead plus Technical Manager",
        "- Broad lead-role use across OEMs, Tier-1s, EV startups, engineering service providers, and automotive software teams",
        "- ATS terms are woven into normal sections rather than placed in a standalone keyword section",
        "",
        "## Version",
        f"- Resume package version: `{VERSION}`",
        "- Generated from Kuldeep Singh's executive resume details and the R01/R02 version-controlled resume history.",
        "",
    ]
    (out_dir / "README.md").write_text("\n".join(readme), encoding="utf-8")


def main():
    for path in [OUTPUT_ROOT, REPO_R03, LOCAL_ROOT]:
        clean_dir(path)

    docx_path, txt_path = build_resume(OUTPUT_ROOT)
    build_readme(OUTPUT_ROOT)

    for file in OUTPUT_ROOT.iterdir():
        if file.is_file():
            shutil.copy2(file, REPO_R03 / file.name)
            shutil.copy2(file, LOCAL_ROOT / file.name)

    print(OUTPUT_ROOT)
    print(docx_path)
    print(txt_path)


if __name__ == "__main__":
    main()
