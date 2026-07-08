from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kuldeepsingh/Documents/Codex/2026-07-08/en")
OUTPUT_ROOT = ROOT / "outputs" / "R02_Targeted_Resume_Pack"
REPO_ROOT = ROOT / "work" / "KuldeepSinghResumeFormat"
REPO_R02 = REPO_ROOT / "resumes" / "R02"

NAME = "Kuldeep Singh"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"
BASE_TITLE = "Technical Manager | EV ECU Software, MATLAB/Simulink MBD & Electrification Systems"

BASE_FACTS = {
    "experience_years": "8 years",
    "company": "IX Energy Pvt. Ltd., Noida, India",
    "dates": "Jul 2018 - Present",
    "education": "B.Tech - Mechanical Engineering, VIT University, Vellore, Tamil Nadu | 2014 - 2018 | First Class, 75%",
    "certs": [
        "Six Sigma Black Belt - Certified",
        "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
        "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
        "AWS IoT & SIMULINK - Amazon Web Services / MathWorks Certified",
    ],
}


TARGETS = [
    {
        "folder": "01_Tata_TataTechnologies_EV_ECU_MBD",
        "doc_label": "Tata_TataTechnologies_EV_ECU_MBD",
        "headline": "EV ECU Software Lead | Tata Commercial EV, eVCU/BMS, MBD & Homologation",
        "target_roles": [
            "EV ECU / eVCU Software Lead",
            "Model-Based Development Lead - Vehicle Control / BMS",
            "Embedded Software / Electrical & Electronics Engineering - Commercial EV",
            "Technical Manager - EV Powertrain Integration",
        ],
        "summary": "Senior EV ECU software and electrification leader with 8 years delivering commercial EV and hybrid powertrain programs across TATA, Mahindra, Daimler Fuso, and government fleet platforms. Strong match for Tata Motors and Tata Technologies roles requiring MATLAB/Simulink/Stateflow model-based development, eVCU/BMS application software, vehicle control logic, DCDC/OBC/PDU integration, V-model delivery, ISO 26262 awareness, supplier coordination, and ICAT/ARAI homologation support. Built and integrated EV systems including the ARAI-homologated TATA 407 LCV/LPT EV conversion with 80 kW powertrain and 53 kWh / 320V LFP battery.",
        "strengths": [
            "Developed BMS and eVCU application software using MATLAB, Simulink, and Stateflow, converting functional and safety requirements into model-based logic for vehicle control, charging coordination, power-up/power-down sequencing, contactor/precharge handling, and HV-LV handshake behaviour.",
            "Defined SOR and component specifications for motor, HV battery, PDU, DCDC, OBC, gearbox, controller, EPS, BCS, TCS, and HVAC systems, aligning supplier deliverables with Tata-style engineering gates and commercial vehicle programme milestones.",
            "Delivered complete EV and P4 hybrid system architecture from concept through validation and ICAT/ARAI homologation, with strong exposure to commercial vehicle integration, quality tools, DFMEA, RCA, and Six Sigma Black Belt methods.",
        ],
        "experience_bullets": [
            "Lead EV and P4 hybrid powertrain architecture, ECU application software, battery engineering, supplier development, vehicle integration, validation, and homologation readiness for commercial vehicle electrification programmes.",
            "Architected EV and P4 hybrid systems from SOR definition through ICAT/ARAI homologation for TATA, Mahindra, and Daimler Fuso platforms, connecting product requirements to hardware, software, testing, and certification outputs.",
            "Implemented eVCU/BMS logic for vehicle operating states, drive enable conditions, charging mode, DCDC/OBC/PDU coordination, diagnostics, fault response, and safe shutdown behaviour using MATLAB/Simulink/Stateflow.",
            "Built EV/hybrid range estimation models and used DFMEA, RCA, DMAIC, SPC, and control charts to improve design robustness, supplier quality, and validation-cycle efficiency.",
            "Collaborated with senior technical leadership at TATA AutoComp Systems on 16T GVW P4 Hybrid ICAT certification and with Bharat Forge/Kalyani Group on 6T GVW EV ARAI certification.",
        ],
        "project_bullets": [
            "TATA 407 LCV/LPT EV conversion: ARAI-homologated 6.5T commercial EV, 80 kW powertrain, 53 kWh / 320V LFP battery, 119 km range; contributed MATLAB/Simulink eVCU/BMS logic, vehicle control, DCDC/PDU/OBC integration, charging interfaces, diagnostics, safety logic, and homologation-readiness support.",
            "TATA 1512 P4 Hybrid bus: ICAT-certified 16T GVW hybrid platform with super-capacitor energy storage and validated 25% fuel-efficiency improvement; supported vehicle integration, pilot fleet trials, validation, and certification readiness.",
            "In-house EV components: developed LV/HV PDU, STM32F4/F1 instrument cluster, Quectel dual-CAN telematics unit, and ultra-capacitor cell monitoring/control system for embedded vehicle electronics programmes.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, Model-Based Design, eVCU, BMS, VCU, Embedded C/C++, STM32, CAN, Vector CANoe, Peak CAN, Bus Master, DBC, ISO 26262, IATF 16949, DFMEA, RCA, Six Sigma, SOR, ARAI, ICAT, commercial EV homologation.",
        "apply_links": [
            ("Tata Motors Careers", "https://www.tatamotors.com/careers/"),
            ("Tata Motors Open Opportunities", "https://careers.tatamotors.com/go/View-All-Opportunities/3505801/"),
            ("Tata Technologies Careers", "https://www.tatatechnologies.com/en/careers/"),
        ],
        "research_notes": [
            "Tata Motors careers route candidates through engineering, R&D, manufacturing, quality, and other opportunity groups; its careers site points applicants to the opportunities portal.",
            "Tata Motors innovation messaging emphasizes electrification, connectivity, digital services, software-on-wheels, and future-ready mobility.",
            "Tata Technologies positions embedded engineering around V-cycle software development, MBD, MIL, SIL, HIL, AUTOSAR, SDV, and automotive engineering services.",
        ],
        "search_terms": "EV ECU, eVCU, BMS Software, Model Based Development, Simulink, Stateflow, Vehicle Control, DCDC, OBC, PDU, Commercial EV, Homologation, ARAI, ICAT, Electrical & Electronics, Embedded Software",
        "sources": [
            "https://www.tatamotors.com/careers/",
            "https://careers.tatamotors.com/go/View-All-Opportunities/3505801/",
            "https://www.tatatechnologies.com/us/solutions/",
        ],
    },
    {
        "folder": "02_Tesla_BMS_Charging_Vehicle_Software",
        "doc_label": "Tesla_BMS_Charging_Vehicle_Software",
        "headline": "BMS, Charging & Vehicle Software Lead | Controls Validation and Firmware Integration",
        "target_roles": [
            "Embedded Firmware Engineer - Battery Management System",
            "Software Integration Engineer - Power Conversion & Charging",
            "Vehicle Software Controls Validation Engineer",
            "Controls / Firmware Engineer - Charging & Energy Products",
        ],
        "summary": "EV software and electrification leader targeting Tesla vehicle software, BMS firmware, charging integration, and controls validation roles. Brings 8 years of safety-critical EV powertrain experience across eVCU/BMS application software, MATLAB/Simulink/Stateflow model-based controls, HV battery architecture, charging controller logic, DCDC/OBC/PDU integration, fault handling, CAN diagnostics, and vehicle-level validation. Delivered certified commercial EV retrofit systems including ARAI-homologated TATA 407 EV work and certified HV battery packs from 5 kWh to 300 kWh across LFP, NMC, LTO, and ultra-capacitor chemistries.",
        "strengths": [
            "Built BMS and eVCU application software for EV systems where battery state, charger state, contactor/precharge conditions, voltage/current limits, and vehicle operating modes must work together safely.",
            "Led AC/DC charging controller logic and power conversion interfaces across charger, BMS, VCU/eVCU, contactors, DCDC, OBC, and PDU systems, matching Tesla-style emphasis on firmware support for charging, power conversion, and highly physical products.",
            "Validated EV/hybrid systems through model checks, bench/vehicle testing, CAN diagnostics, fault analysis, calibration support, and certification documentation, with DFMEA/RCA and ISO 26262 functional-safety awareness.",
        ],
        "experience_bullets": [
            "Developed BMS and eVCU software using MATLAB/Simulink/Stateflow MBD methodology, translating functional and safety requirements into control logic for drive readiness, charging mode, fault mode, and safe shutdown.",
            "Implemented charging controller and power conversion logic across AC/DC charging architectures, DCDC converter, OBC, PDU, HV battery, contactors, and vehicle safety interfaces.",
            "Designed HV battery systems from 5 kWh to 300 kWh across 72V-800V LFP, NMC, LTO, and ultra-capacitor architectures, combining software behaviour with pack design, BMU development, and BTMS/HVAC validation.",
            "Supported vehicle-level software integration and validation using CAN tools, diagnostics, issue debugging, calibration behaviour review, and structured problem solving.",
            "Managed cross-functional suppliers and internal teams across battery, motor, gearbox, DCDC, EPS, BCS, TCS, and HVAC to maintain component readiness and release-quality integration outputs.",
        ],
        "project_bullets": [
            "TATA 407 EV conversion: ARAI-homologated 6.5T EV with 80 kW drive and 53 kWh / 320V LFP battery; contributed eVCU/BMS logic, charging interfaces, DCDC/OBC/PDU coordination, diagnostics, and safe operating-state control.",
            "Certified battery packs: delivered 53 kWh/332V LFP electric truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC pack, and 0.5 kWh/400V ultra-capacitor hybrid bus pack.",
            "EV charging controller: led AC/DC charging architecture logic, HV-LV handshake, charger/BMS/VCU coordination, fault response, and vehicle safety readiness.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, Embedded C/C++, BMS, eVCU, HV battery, power conversion, charging firmware, AC/DC charging, DCDC, OBC, PDU, CAN, diagnostics, Vector CANoe, Peak CAN, ISO 26262, DFMEA, RCA, validation, calibration.",
        "apply_links": [
            ("Tesla Careers Search", "https://www.tesla.com/careers/search"),
            ("Tesla - Embedded Firmware Engineer, BMS", "https://www.tesla.com/careers/search/job/embedded-firmware-engineer-battery-management-system-254793"),
            ("Tesla - Power Conversion & Charging Integration", "https://www.tesla.com/careers/search/job/software-integration-engineer-power-conversion-charging-255576"),
            ("Tesla - Vehicle Software Controls Validation", "https://www.tesla.com/careers/search/job/software-engineer-controls-systems-validation-vehicle-software--237400"),
        ],
        "research_notes": [
            "Current Tesla examples include BMS embedded firmware, power conversion and charging integration, and vehicle software controls validation roles.",
            "The strongest resume angle is hands-on BMS/eVCU software, charging control, power conversion, validation, and production-quality EV system integration.",
            "Tesla roles are very software-and-validation specific, so the resume should foreground firmware logic, model-based controls, diagnostics, fault response, and vehicle-level testing rather than broad management first.",
        ],
        "search_terms": "BMS Firmware, Charging Firmware, Power Conversion, Vehicle Software, Controls Validation, Battery Management System, DCDC, OBC, HV Battery, CAN Diagnostics, Embedded C, Simulink, Stateflow, Vehicle Integration",
        "sources": [
            "https://www.tesla.com/careers/search",
            "https://www.tesla.com/careers/search/job/embedded-firmware-engineer-battery-management-system-254793",
            "https://www.tesla.com/careers/search/job/software-integration-engineer-power-conversion-charging-255576",
            "https://www.tesla.com/careers/search/job/software-engineer-controls-systems-validation-vehicle-software--237400",
        ],
    },
    {
        "folder": "03_Daimler_Mercedes_Commercial_EV_SDV",
        "doc_label": "Daimler_Mercedes_Commercial_EV_SDV",
        "headline": "Commercial EV Software & Electronics Lead | Daimler Truck / Mercedes-Benz SDV",
        "target_roles": [
            "Embedded Software Architect / Engineer - Commercial Vehicles",
            "EV Powertrain Software Lead - Trucks & Buses",
            "Software and Electronics Engineer - DTICI / MBRDI",
            "Vehicle Systems Integration Lead - Electrification",
        ],
        "summary": "Commercial EV and hybrid powertrain software leader with 8 years delivering truck, bus, and retrofit electrification systems from architecture through validation and ICAT/ARAI homologation. Targeting Daimler Truck, DTICI, and Mercedes-Benz engineering roles requiring embedded automotive software, MATLAB/Simulink, CAN/LIN/Ethernet mindset, AUTOSAR basics, ISO 26262, software-defined vehicle readiness, supplier coordination, and global team collaboration. Directly supported Daimler Fuso 5T EV architecture and collaborated with senior OEM/Tier-1 leaders on certified commercial EV and P4 hybrid programmes.",
        "strengths": [
            "Strong commercial vehicle fit: TATA 407 EV, Daimler Fuso 5T EV, TATA 1512 P4 Hybrid bus, and 6T GVW EV ARAI certification exposure across trucks and buses.",
            "Developed eVCU/BMS application software and system logic for power-up/power-down sequencing, vehicle operating states, DCDC/OBC/PDU integration, charging controller interfaces, diagnostics, and HV-LV handshakes.",
            "Comfortable with Daimler-style software/electronics requirements: MATLAB, CAN, embedded software, functional safety, AUTOSAR basics, system decomposition, global coordination, supplier quality, and change-ready product development.",
        ],
        "experience_bullets": [
            "Led EV and P4 hybrid powertrain architecture, ECU software, battery engineering, supplier development, validation, and homologation readiness for commercial vehicle platforms.",
            "Supported Daimler Fuso 5T EV conversion architecture with 80 kW powertrain, 56 kWh / 310V NMC battery system, power electronics interfaces, HV battery integration, vehicle validation, and supplier coordination.",
            "Developed MATLAB/Simulink/Stateflow eVCU and BMS logic for functional requirements, safety requirements, V-model verification artefacts, and ISO 26262 ASIL-aligned control behaviour.",
            "Integrated vehicle communication and software interfaces across CAN-connected battery, charger, DCDC, OBC, motor controller, PDU, telematics, instrument cluster, EPS, BCS, TCS, and HVAC systems.",
            "Managed cross-functional teams and global suppliers while applying DFMEA, RCA, Six Sigma, and validation planning to close integration issues and improve release readiness.",
        ],
        "project_bullets": [
            "Daimler Fuso 5T EV: supported EV conversion architecture, 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, power electronics interfaces, and vehicle validation.",
            "TATA 407 EV and 6T GVW EV ARAI programme: delivered commercial EV software/system integration exposure directly aligned with electric truck and bus development.",
            "TATA 1512 P4 Hybrid ICAT programme: collaborated with TATA AutoComp senior leadership on 16T GVW hybrid bus certification and state government pilot fleet trials.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, Embedded C/C++, CAN, DBC, Vector CANoe, Peak CAN, AUTOSAR basics, ISO 26262, V-model, Git awareness, eVCU, BMS, DCDC, OBC, PDU, HV battery, commercial EV, SDV, trucks, buses, homologation.",
        "apply_links": [
            ("DTICI Careers", "https://dtici.daimlertruck.com/career/"),
            ("Daimler Truck Job Search", "https://jobsearch.daimlertruck.com/"),
            ("Daimler Truck Innovation Center India - Bengaluru", "https://locations.daimlertruck.com/en/career/locations/detail/bengaluru-daimler-truck-innovation-center-india-private-limited"),
            ("Mercedes-Benz Careers", "https://group.mercedes-benz.com/careers/"),
        ],
        "research_notes": [
            "DTICI careers highlight Product Engineering, Software and Electronics, and IT, looking for technical skills, problem solving, teamwork, innovation, and growth mindset.",
            "Daimler Truck job examples emphasize automotive software development, MATLAB, Vector tools, embedded software, AUTOSAR, CAN/LIN/Ethernet, ISO 26262, software-defined vehicles, and global team coordination.",
            "DTICI Bengaluru messaging specifically mentions software tools, hardware labs, connectivity, autonomous, cybersecurity, system integration labs, and electrification topics.",
        ],
        "search_terms": "Daimler Truck, DTICI, Mercedes-Benz, Embedded Software, Software and Electronics, Commercial Vehicle, Electric Truck, MATLAB, CAN, AUTOSAR, ISO 26262, SDV, VCU, BMS, DCDC, OBC",
        "sources": [
            "https://dtici.daimlertruck.com/career/",
            "https://jobsearch.daimlertruck.com/index.php?ac=jobad&id=422397",
            "https://locations.daimlertruck.com/en/career/locations/detail/bengaluru-daimler-truck-innovation-center-india-private-limited",
            "https://group.mercedes-benz.com/careers/",
        ],
    },
    {
        "folder": "04_BMW_Embedded_ECU_Software",
        "doc_label": "BMW_Embedded_ECU_Software",
        "headline": "Embedded ECU Software Lead | BMW Car IT, Vehicle Controls & MBD",
        "target_roles": [
            "Embedded Software Developer - ECU Platform",
            "Automotive Software Engineer - Vehicle Controls",
            "Software Integration / Systems Engineer - Electronic Control Units",
            "Model-Based Development Engineer - BMW Group",
        ],
        "summary": "Automotive embedded software and EV systems leader targeting BMW Group and BMW Car IT roles in ECU software, vehicle controls, software integration, and model-based development. Offers 8 years of hands-on eVCU/BMS application software, MATLAB/Simulink/Stateflow controls, embedded C/C++, CAN diagnostics, HV battery integration, and safety-critical EV powertrain delivery. Brings proven system-level experience from commercial EV conversions, P4 hybrid buses, certified HV battery packs, and vehicle electronics including instrument clusters, dual-CAN telematics, and LV/HV PDU development.",
        "strengths": [
            "BMW Car IT alignment: practical experience with electronic control unit software, related embedded tools, system decomposition, and vehicle-level software integration.",
            "Developed model-based vehicle control and BMS software across drive readiness, charging, DCDC/OBC/PDU integration, fault handling, diagnostics, and safe operating states.",
            "Delivered embedded vehicle electronics including STM32F4/F1 instrument cluster, Quectel dual-CAN telematics, ultra-capacitor monitoring, and LV/HV power distribution control.",
        ],
        "experience_bullets": [
            "Developed BMS and eVCU application software with MATLAB/Simulink/Stateflow MBD methodology, aligning functional requirements, safety requirements, and verification outputs with V-model discipline.",
            "Implemented vehicle control logic, power-up/power-down sequencing, DCDC/OBC/PDU coordination, charging controller interfaces, HV-LV handshake protocols, and diagnostics for EV and hybrid platforms.",
            "Built component-level and vehicle-level range estimation models and integrated battery, motor, gearbox, controller, EPS, BCS, TCS, and HVAC systems through structured interface definition.",
            "Developed STM32-based instrument cluster, dual-CAN telematics, LV/HV PDU, and ultra-capacitor monitoring/control system, combining embedded software with vehicle electronics requirements.",
            "Applied ISO 26262 awareness, IATF 16949 process thinking, DFMEA, RCA, and Six Sigma methods to improve design robustness and validation efficiency.",
        ],
        "project_bullets": [
            "TATA 407 EV: ARAI-homologated commercial EV with 80 kW drive and 53 kWh LFP battery; contributed ECU software, eVCU/BMS logic, charging, diagnostics, DCDC/OBC/PDU coordination, and vehicle control.",
            "In-house embedded vehicle electronics: built STM32F4/F1 instrument cluster, Quectel dual-CAN telematics, LV/HV PDU, and ultra-capacitor cell monitoring/control system.",
            "HV battery systems: designed and packaged 5 kWh to 300 kWh battery systems across 72V-800V LFP, NMC, LTO, and ultra-capacitor architectures.",
        ],
        "tools": "Embedded C/C++, MATLAB, Simulink, Stateflow, ECU software, eVCU, BMS, VCU, STM32, NXP, Infineon, CAN, DBC, Vector CANoe, Peak CAN, diagnostics, ISO 26262, AUTOSAR basics, V-model, EV powertrain, HV battery.",
        "apply_links": [
            ("BMW Group Careers", "https://www.bmwgroup.jobs/"),
            ("BMW Group IT Careers", "https://www.bmwgroup.jobs/us/en/jobfields/information-technology.html"),
            ("BMW Car IT Careers", "https://www.bmw-carit.de/en/career.html"),
        ],
        "research_notes": [
            "BMW Car IT says it designs and implements components for BMW electronic control units and related tools/frameworks.",
            "BMW Group IT messaging links software developers with automotive IT, HMI, driver assistance, and shaping future vehicles.",
            "For BMW, keep the resume language precise: ECU software, embedded C/C++, systems integration, vehicle electronics, CAN diagnostics, model-based controls, and product-quality delivery.",
        ],
        "search_terms": "BMW Car IT, BMW Group, Embedded Software, ECU Platform, Electronic Control Units, Automotive IT, Vehicle Controls, Simulink, Stateflow, Embedded C++, CAN, ISO 26262, AUTOSAR, BMS, eVCU",
        "sources": [
            "https://www.bmw-carit.de/en/career.html",
            "https://www.bmwgroup.jobs/us/en/jobfields/information-technology.html",
            "https://www.bmwgroup.jobs/",
        ],
    },
    {
        "folder": "05_JLR_LandRover_EDU_EV_Validation",
        "doc_label": "JLR_LandRover_EDU_EV_Validation",
        "headline": "EV Powertrain Validation & Controls Lead | JLR / Land Rover EDU, Battery & Software",
        "target_roles": [
            "EV Powertrain Validation Engineer",
            "Electric Drive Unit / EDU Validation Engineer",
            "Vehicle Controls & Systems Integration Engineer",
            "Battery / Powertrain Software Validation Lead",
        ],
        "summary": "EV powertrain software and validation leader targeting JLR / Land Rover roles across Electric Drive Unit validation, powertrain testing, vehicle controls, battery integration, and software-enabled test/data workflows. Brings 8 years in EV and hybrid commercial platforms with MATLAB/Simulink/Stateflow controls, BMS/eVCU software, HV battery architecture, DCDC/OBC/PDU integration, CAN diagnostics, test planning, DFMEA/RCA, Six Sigma, and ICAT/ARAI homologation. Strong fit for JLR roles requiring powertrain testing at component, subsystem, and system levels plus cross-functional issue resolution.",
        "strengths": [
            "Powertrain validation fit: delivered EV and hybrid platforms from architecture to vehicle validation, certification sign-off support, and issue closure using structured problem-solving tools.",
            "MATLAB/Simulink and data workflow fit: built range estimation models, control strategies, and verification artefacts; comfortable using model/data analysis to support efficient validation decisions.",
            "EDU/electrification fit: integrated motor, gearbox, HV battery, DCDC, OBC, PDU, charger, BMS, controller, and vehicle controls across commercial EV platforms.",
        ],
        "experience_bullets": [
            "Led EV and hybrid system integration across battery, motor, gearbox, DCDC, OBC, PDU, charger, controllers, sensors, telematics, and vehicle control interfaces.",
            "Developed eVCU/BMS application software and operating-state logic using MATLAB/Simulink/Stateflow for drive readiness, charging mode, fault handling, power sequencing, and safe shutdown behaviour.",
            "Supported vehicle-level validation, test planning, fault investigation, CAN diagnostics, calibration behaviour review, supplier issue closure, and homologation evidence preparation.",
            "Applied DFMEA, RCA, DMAIC, Six Sigma, and quality tools to resolve integration issues, reduce validation cycles, and improve design robustness.",
            "Managed EV programme deliverables across internal engineering teams and external suppliers, communicating technical risk and readiness at system, subsystem, and component levels.",
        ],
        "project_bullets": [
            "TATA 407 EV: ARAI-homologated 6.5T EV with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed ECU logic, charging interfaces, HV battery coordination, diagnostics, and validation support.",
            "Daimler Fuso 5T EV: supported 80 kW EV conversion with 56 kWh / 310V NMC pack, power electronics interfaces, vehicle validation, and supplier coordination.",
            "P4 Hybrid Bus: integrated super-capacitor electrification kit for TATA 1512, supported ICAT validation, and contributed to 25% fuel-efficiency improvement evidence.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, powertrain validation, Electric Drive Unit, EDU, HV battery, BMS, eVCU, CAN, Vector CANoe, Peak CAN, test planning, instrumentation, DFMEA, RCA, Six Sigma, vehicle integration, ISO 26262, ICAT, ARAI.",
        "apply_links": [
            ("JLR Careers Search", "https://www.jaguarlandrovercareers.com/content/Find-a-job/?locale=en_GB"),
            ("JLR Careers Home", "https://www.jaguarlandrovercareers.com/"),
            ("JLR EDU Validation Example Role", "https://www.jaguarlandrovercareers.com/job/Budapest-DV-Engineer-EDU-Validation/1365486733/"),
        ],
        "research_notes": [
            "JLR careers emphasizes all-electric future, engineering, technology, and global R&D investment.",
            "A current EDU validation role highlights component/subsystem/system powertrain testing, rig/test equipment, MATLAB/Python/data post-processing, quality tools, issue resolution, and cross-functional communication.",
            "Use JLR-facing language around validation, DVP, test evidence, EDU, powertrain, issue resolution, and system-level engineering sign-off.",
        ],
        "search_terms": "JLR, Land Rover, Electric Drive Unit, EDU, Powertrain Validation, EV Controls, Battery Integration, MATLAB, INCA, Testlab, CAN, DVP, DFMEA, Six Sigma, Vehicle Integration, ARAI, ICAT",
        "sources": [
            "https://www.jaguarlandrovercareers.com/",
            "https://www.jaguarlandrovercareers.com/content/Find-a-job/?locale=en_GB",
            "https://www.jaguarlandrovercareers.com/job/Budapest-DV-Engineer-EDU-Validation/1365486733/",
        ],
    },
    {
        "folder": "06_AVL_EMobility_Simulation_Test_Systems",
        "doc_label": "AVL_EMobility_Simulation_Test_Systems",
        "headline": "E-Mobility, Simulation & Test Systems Lead | AVL-Style EV Software Engineering",
        "target_roles": [
            "E-Mobility Systems Engineer",
            "Software Development Engineer - EV Test / Simulation",
            "Battery / E-Drive Test Systems Engineer",
            "Model-Based Development / System Integration Engineer",
        ],
        "summary": "E-mobility software and systems leader targeting AVL roles in electrification, simulation, test/validation systems, charging solutions, and vehicle efficiency engineering. Offers 8 years delivering EV/hybrid systems with MATLAB/Simulink/Stateflow MBD, range estimation models, BMS/eVCU software, charging controller logic, HV battery systems, DCDC/OBC/PDU integration, and validation/homologation support. Strongly aligned with AVL’s focus on engineering, simulation, testing, electrification, software, AI/automation, battery/e-drive/inverter validation, and model-based virtual development before hardware build.",
        "strengths": [
            "Built MBD models for component-level and vehicle-level EV/hybrid range estimation, supporting virtual design decisions, vehicle efficiency studies, and validation planning.",
            "Led EV charging controller logic and system integration across charger, BMS, eVCU/VCU, DCDC, OBC, PDU, HV battery, contactors, and vehicle safety states.",
            "Designed HV battery systems and certified battery packs across LFP, NMC, LTO, and ultra-capacitor chemistries, matching AVL e-mobility work around batteries, e-motors, inverters, simulation, testing, and charging solutions.",
        ],
        "experience_bullets": [
            "Developed MATLAB/Simulink/Stateflow control models for eVCU/BMS functions, range estimation, charging logic, and EV/hybrid operating-state behaviour.",
            "Integrated EV components and test/validation workflows across battery, motor, gearbox, DCDC, OBC, PDU, charger, controller, telematics, and thermal systems.",
            "Applied DFMEA, RCA, Six Sigma, and validation evidence reviews to improve design robustness and shorten test/debug cycles across EV programmes.",
            "Led supplier coordination and hardware-software interface definition for battery, motor, gearbox, DCDC, EPS, BCS, TCS, HVAC, and controller systems.",
            "Supported vehicle-level validation, certification documentation, and homologation readiness for ICAT/ARAI-certified EV and hybrid programmes.",
        ],
        "project_bullets": [
            "E-mobility models: created EV/hybrid range estimation and control models for component and vehicle-level scenarios using MATLAB/Simulink.",
            "Battery and charging systems: delivered 5 kWh to 300 kWh HV battery systems, AC/DC charging controller logic, DCDC/OBC/PDU integration, and BTMS/HVAC validation.",
            "Commercial EV programmes: supported TATA 407 EV, Daimler Fuso 5T EV, TATA 1512 P4 Hybrid, Ambassador, and Mahindra Supro EV conversion programmes through integration and validation.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, model-based simulation, range estimation, e-mobility, EV test systems, BMS, eVCU, charging systems, DCDC, OBC, PDU, battery testing, e-drive, inverter, CAN, diagnostics, DFMEA, RCA, Six Sigma.",
        "apply_links": [
            ("AVL Careers", "https://www.avl.com/en/career"),
            ("AVL E-Mobility Careers", "https://www.avl.com/en/career/fields-interest/e-mobility"),
            ("AVL IT and Software Careers", "https://www.avl.com/en/career/fields-interest/it-and-software"),
            ("AVL Jobs Search", "https://jobs.avl.com/"),
        ],
        "research_notes": [
            "AVL’s e-mobility careers page calls out development support, test and validation solutions, simulation tools, system integration, and charging-system solutions.",
            "AVL’s IT/software page emphasizes driving attribute simulation, vehicle efficiency, virtual vehicle characteristic engineering, and designing/testing vehicle concepts before hardware build.",
            "AVL corporate messaging emphasizes engineering, simulation, testing, electrification, software, AI, automation, software-defined vehicles, vehicle software, embedded systems, and battery-electric/fuel-cell systems.",
        ],
        "search_terms": "AVL, e-mobility, simulation, test systems, vehicle efficiency, charging systems, BMS, e-drive, inverter, battery validation, MATLAB, Simulink, Stateflow, system integration, virtual vehicle",
        "sources": [
            "https://www.avl.com/en/career",
            "https://www.avl.com/en/career/fields-interest/e-mobility",
            "https://www.avl.com/en/career/fields-interest/it-and-software",
            "https://www.avl.com/en/about-avl",
        ],
    },
    {
        "folder": "07_Bosch_Automotive_Embedded_Systems",
        "doc_label": "Bosch_Automotive_Embedded_Systems",
        "headline": "Automotive Embedded Systems & EV Integration Lead | Bosch Mobility Fit",
        "target_roles": [
            "System Integration Engineer - Automotive Applications",
            "Embedded Software Engineer - Automotive Products",
            "EV / Cross-Domain Vehicle Systems Engineer",
            "BMS / ECU Software Integration Lead",
        ],
        "summary": "Automotive embedded systems and EV software leader targeting Bosch Mobility roles involving cross-domain automotive products, embedded software, hardware electronics integration, future vehicle architectures, and system integration. Brings 8 years of eVCU/BMS software, MATLAB/Simulink/Stateflow MBD, embedded C/C++, STM32/NXP/Infineon exposure, CAN diagnostics, HV battery systems, DCDC/OBC/PDU integration, charging controller logic, and vehicle-level validation. Strong fit for Bosch roles connecting hardware electronics, embedded software, software architecture, automotive product development, and global engineering delivery.",
        "strengths": [
            "Bosch integration fit: developed cross-domain EV systems linking hardware electronics, embedded software, CAN communication, power electronics, battery systems, vehicle controls, and validation.",
            "Embedded systems fit: developed STM32-based instrument cluster, dual-CAN telematics, ultra-capacitor monitoring/control, LV/HV PDU, and eVCU/BMS application software.",
            "Automotive product fit: delivered SOR, supplier coordination, validation evidence, quality tools, and ICAT/ARAI homologation readiness for commercial EV and hybrid programmes.",
        ],
        "experience_bullets": [
            "Integrated hardware electronics and embedded software across HV battery, BMS/eVCU, DCDC, OBC, PDU, motor controller, telematics, sensors, and vehicle controls.",
            "Developed MATLAB/Simulink/Stateflow eVCU and BMS application software for functional requirements, safety requirements, diagnostics, fault handling, and safe vehicle operating states.",
            "Built STM32F4/F1 instrument cluster, Quectel dual-CAN telematics, LV/HV PDU, and ultra-capacitor cell monitoring/control system for in-house EV component development.",
            "Defined component/system specifications and SOR documents for supplier sourcing and integration across battery, motor, gearbox, DCDC, EPS, BCS, TCS, HVAC, and controller systems.",
            "Used CAN tools, diagnostics, DFMEA, RCA, Six Sigma, ISO 26262 awareness, and IATF 16949 process thinking to support robust automotive product development.",
        ],
        "project_bullets": [
            "TATA 407 EV: ARAI-homologated commercial EV with eVCU/BMS logic, vehicle control, charging interfaces, DCDC/OBC/PDU integration, diagnostics, and HV battery coordination.",
            "Embedded component suite: delivered LV/HV PDU, STM32 instrument cluster, dual-CAN telematics, ultra-capacitor monitoring/control, and battery-related software interfaces.",
            "HV battery systems: packaged and integrated 5 kWh to 300 kWh packs across LFP, NMC, LTO, and ultra-capacitor technologies with BTMS/HVAC validation.",
        ],
        "tools": "Embedded C/C++, STM32, Microchip PIC, NXP, Infineon, MATLAB, Simulink, Stateflow, BMS, eVCU, CAN, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, DBC, diagnostics, software architecture, system integration, ISO 26262, IATF 16949.",
        "apply_links": [
            ("Bosch Global Careers", "https://www.bosch.com/careers/"),
            ("Bosch Jobs Search", "https://jobs.bosch.com/"),
            ("Bosch System Integration Example Role", "https://jobs.bosch.com/it/job/REF288304D-system-integration-engineer-for-automotive-applications-communication-specialist"),
        ],
        "research_notes": [
            "Bosch careers emphasizes professional opportunities across fields of work and worldwide locations.",
            "A Bosch automotive systems integration role focuses on integrating hardware electronics and embedded software for next-generation software architectures across vehicle types.",
            "For Bosch, lead with embedded systems, cross-domain integration, hardware/software interface ownership, automotive software architecture, CAN, BMS/eVCU, and supplier coordination.",
        ],
        "search_terms": "Bosch Mobility, Embedded Software, System Integration, Automotive Applications, Hardware Electronics, Software Architecture, Future Vehicles, BMS, eVCU, CAN, STM32, DCDC, OBC, PDU, ISO 26262",
        "sources": [
            "https://www.bosch.com/careers/",
            "https://jobs.bosch.com/",
            "https://jobs.bosch.com/it/job/REF288304D-system-integration-engineer-for-automotive-applications-communication-specialist",
        ],
    },
    {
        "folder": "08_Continental_Software_System_Engineering",
        "doc_label": "Continental_Software_System_Engineering",
        "headline": "Software & System Engineering Lead | EV Controls, BMS and Vehicle Integration",
        "target_roles": [
            "Software Engineering / System Engineering - Automotive",
            "EV Systems Integration Engineer",
            "Embedded Software Engineer - Mobility Systems",
            "BMS / Vehicle Controls Lead",
        ],
        "summary": "Automotive software and systems engineering leader targeting Continental roles in software engineering, system engineering, safer mobility, sustainable vehicle technologies, and future mobility solutions. Brings 8 years across EV/hybrid architecture, eVCU/BMS software, MATLAB/Simulink/Stateflow MBD, embedded C/C++, CAN diagnostics, HV battery systems, DCDC/OBC/PDU integration, vehicle validation, and quality-led development. Strong fit for Continental roles needing cross-functional systems thinking, embedded software, EV powertrain integration, safety/process discipline, and customer-ready mobility products.",
        "strengths": [
            "Software/system engineering fit: led EV architecture, ECU application software, component specifications, supplier interfaces, validation planning, and homologation readiness.",
            "Mobility technology fit: delivered EV and hybrid systems that improve efficiency, safety, and sustainability across commercial vehicle applications.",
            "Process fit: combined ISO 26262 functional safety awareness, IATF 16949 quality management, DFMEA, RCA, Six Sigma, and validation evidence into structured automotive delivery.",
        ],
        "experience_bullets": [
            "Led EV and P4 hybrid architecture, ECU application software, HV battery engineering, supplier development, validation, and certification readiness for commercial vehicle platforms.",
            "Developed MATLAB/Simulink/Stateflow eVCU/BMS logic for vehicle control, charging, DCDC/OBC/PDU coordination, fault handling, diagnostics, and safe operating states.",
            "Defined system and component specifications for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, EPS, BCS, TCS, HVAC, and embedded vehicle electronics.",
            "Managed cross-functional engineering and supplier teams to improve component readiness, validation efficiency, system integration, and release-quality software/hardware interfaces.",
            "Applied DFMEA, RCA, DMAIC, SPC, and control charts to reduce risk, improve supplier quality, and support automotive compliance deliverables.",
        ],
        "project_bullets": [
            "Commercial EV systems: TATA 407 EV, Daimler Fuso 5T EV, Ambassador, and Mahindra Supro EV conversions with HV battery, charger, DCDC, controller, vehicle control, and validation integration.",
            "P4 Hybrid Bus: TATA 1512 ICAT-certified hybrid programme with super-capacitor pack and 25% fuel-efficiency improvement validation.",
            "In-house software/electronics: LV/HV PDU, STM32 instrument cluster, dual-CAN telematics, ultra-capacitor cell monitoring/control, and EV range estimation models.",
        ],
        "tools": "Software Engineering, System Engineering, MATLAB, Simulink, Stateflow, Embedded C/C++, BMS, eVCU, VCU, CAN, DBC, diagnostics, HV battery, DCDC, OBC, PDU, vehicle integration, ISO 26262, IATF 16949, DFMEA, RCA, Six Sigma.",
        "apply_links": [
            ("Continental Careers", "https://www.continental.com/en/career/"),
            ("Continental Career Level - Experienced Professionals", "https://www.continental.com/en/career/your-career-level/experienced-professionals/"),
            ("Continental Job Search", "https://www.continental.com/en/career/"),
        ],
        "research_notes": [
            "Continental careers lists Software Engineering and System Engineering as a main range of jobs and points experienced professionals to opportunities.",
            "Continental career messaging emphasizes tomorrow’s mobility, safety, comfort, and sustainability.",
            "For Continental, the strongest angle is software/system engineering with evidence of EV integration, quality methods, functional safety, supplier coordination, and mobility product delivery.",
        ],
        "search_terms": "Continental, Software Engineering, System Engineering, EV Systems, Embedded Software, BMS, eVCU, CAN, DCDC, OBC, PDU, Vehicle Integration, ISO 26262, IATF 16949, Functional Safety, Mobility",
        "sources": [
            "https://www.continental.com/en/career/",
            "https://www.continental.com/en/career/your-career-level/experienced-professionals/",
        ],
    },
]

FORBIDDEN_RESUME_TERMS = [
    "tata",
    "tesla",
    "daimler",
    "mercedes",
    "bmw",
    "jlr",
    "land rover",
    "landrover",
    "avl",
    "bosch",
    "continental",
]

RESUME_OVERRIDES = {
    "01_Tata_TataTechnologies_EV_ECU_MBD": {
        "resume_label": "EV_ECU_MBD_Commercial_EV",
        "headline": "EV ECU Software Lead | Commercial EV, eVCU/BMS, MBD & Homologation",
        "target_roles": [
            "EV ECU / eVCU Software Lead",
            "Model-Based Development Lead - Vehicle Control / BMS",
            "Embedded Software / Electrical & Electronics Engineer - Commercial EV",
            "Technical Manager - EV Powertrain Integration",
        ],
        "summary": "Senior EV ECU software and electrification leader with 8 years delivering commercial EV and hybrid powertrain programmes across trucks, buses, retrofit EVs, and government fleet platforms. Strong fit for roles requiring MATLAB/Simulink/Stateflow model-based development, eVCU/BMS application software, vehicle control logic, DCDC/OBC/PDU integration, V-model delivery, ISO 26262 awareness, supplier coordination, and ICAT/ARAI homologation support. Built and integrated ARAI-homologated LCV/LPT commercial EV systems with 80 kW powertrain and 53 kWh / 320V LFP battery architecture.",
        "strengths": [
            "Developed BMS and eVCU application software using MATLAB, Simulink, and Stateflow, converting functional and safety requirements into model-based logic for vehicle control, charging coordination, power-up/power-down sequencing, contactor/precharge handling, and HV-LV handshake behaviour.",
            "Defined SOR and component specifications for motor, HV battery, PDU, DCDC, OBC, gearbox, controller, EPS, BCS, TCS, and HVAC systems, aligning supplier deliverables with commercial vehicle engineering gates and programme milestones.",
            "Delivered complete EV and P4 hybrid system architecture from concept through validation and ICAT/ARAI homologation, with strong exposure to commercial vehicle integration, quality tools, DFMEA, RCA, and Six Sigma Black Belt methods.",
        ],
        "experience_bullets": [
            "Lead EV and P4 hybrid powertrain architecture, ECU application software, battery engineering, supplier development, vehicle integration, validation, and homologation readiness for commercial vehicle electrification programmes.",
            "Architected EV and P4 hybrid systems from SOR definition through ICAT/ARAI homologation, connecting product requirements to hardware, software, testing, and certification outputs.",
            "Implemented eVCU/BMS logic for vehicle operating states, drive enable conditions, charging mode, DCDC/OBC/PDU coordination, diagnostics, fault response, and safe shutdown behaviour using MATLAB/Simulink/Stateflow.",
            "Built EV/hybrid range estimation models and used DFMEA, RCA, DMAIC, SPC, and control charts to improve design robustness, supplier quality, and validation-cycle efficiency.",
            "Collaborated with senior OEM and Tier-1 technical leadership on 16T GVW P4 hybrid ICAT certification and 6T GVW EV ARAI certification readiness.",
        ],
        "project_bullets": [
            "LCV/LPT EV conversion: ARAI-homologated 6.5T commercial EV, 80 kW powertrain, 53 kWh / 320V LFP battery, 119 km range; contributed MATLAB/Simulink eVCU/BMS logic, vehicle control, DCDC/PDU/OBC integration, charging interfaces, diagnostics, safety logic, and homologation-readiness support.",
            "16T P4 hybrid bus: ICAT-certified hybrid platform with super-capacitor energy storage and validated 25% fuel-efficiency improvement; supported vehicle integration, pilot fleet trials, validation, and certification readiness.",
            "In-house EV components: developed LV/HV PDU, STM32F4/F1 instrument cluster, Quectel dual-CAN telematics unit, and ultra-capacitor cell monitoring/control system for embedded vehicle electronics programmes.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, Model-Based Design, eVCU, BMS, VCU, Embedded C/C++, STM32, CAN, Vector CANoe, Peak CAN, Bus Master, DBC, ISO 26262, IATF 16949, DFMEA, RCA, Six Sigma, SOR, ARAI, ICAT, commercial EV homologation.",
    },
    "02_Tesla_BMS_Charging_Vehicle_Software": {
        "resume_label": "BMS_Charging_Vehicle_Software",
        "headline": "BMS, Charging & Vehicle Software Lead | Controls Validation and Firmware Integration",
        "target_roles": [
            "Embedded Firmware Engineer - Battery Management System",
            "Software Integration Engineer - Power Conversion & Charging",
            "Vehicle Software Controls Validation Engineer",
            "Controls / Firmware Engineer - Charging & Energy Products",
        ],
        "summary": "EV software and electrification leader targeting vehicle software, BMS firmware, charging integration, and controls validation roles. Brings 8 years of safety-critical EV powertrain experience across eVCU/BMS application software, MATLAB/Simulink/Stateflow model-based controls, HV battery architecture, charging controller logic, DCDC/OBC/PDU integration, fault handling, CAN diagnostics, and vehicle-level validation. Delivered certified commercial EV retrofit systems and certified HV battery packs from 5 kWh to 300 kWh across LFP, NMC, LTO, and ultra-capacitor chemistries.",
        "strengths": [
            "Built BMS and eVCU application software for EV systems where battery state, charger state, contactor/precharge conditions, voltage/current limits, and vehicle operating modes must work together safely.",
            "Led AC/DC charging controller logic and power conversion interfaces across charger, BMS, VCU/eVCU, contactors, DCDC, OBC, and PDU systems, supporting firmware behaviour for charging, power conversion, and highly physical vehicle products.",
            "Validated EV/hybrid systems through model checks, bench/vehicle testing, CAN diagnostics, fault analysis, calibration support, and certification documentation, with DFMEA/RCA and ISO 26262 functional-safety awareness.",
        ],
        "experience_bullets": [
            "Developed BMS and eVCU software using MATLAB/Simulink/Stateflow MBD methodology, translating functional and safety requirements into control logic for drive readiness, charging mode, fault mode, and safe shutdown.",
            "Implemented charging controller and power conversion logic across AC/DC charging architectures, DCDC converter, OBC, PDU, HV battery, contactors, and vehicle safety interfaces.",
            "Designed HV battery systems from 5 kWh to 300 kWh across 72V-800V LFP, NMC, LTO, and ultra-capacitor architectures, combining software behaviour with pack design, BMU development, and BTMS/HVAC validation.",
            "Supported vehicle-level software integration and validation using CAN tools, diagnostics, issue debugging, calibration behaviour review, and structured problem solving.",
            "Managed cross-functional suppliers and internal teams across battery, motor, gearbox, DCDC, EPS, BCS, TCS, and HVAC to maintain component readiness and release-quality integration outputs.",
        ],
        "project_bullets": [
            "LCV/LPT EV conversion: ARAI-homologated 6.5T EV with 80 kW drive and 53 kWh / 320V LFP battery; contributed eVCU/BMS logic, charging interfaces, DCDC/OBC/PDU coordination, diagnostics, and safe operating-state control.",
            "Certified battery packs: delivered 53 kWh/332V LFP electric truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC pack, and 0.5 kWh/400V ultra-capacitor hybrid bus pack.",
            "EV charging controller: led AC/DC charging architecture logic, HV-LV handshake, charger/BMS/VCU coordination, fault response, and vehicle safety readiness.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, Embedded C/C++, BMS, eVCU, HV battery, power conversion, charging firmware, AC/DC charging, DCDC, OBC, PDU, CAN, diagnostics, Vector CANoe, Peak CAN, ISO 26262, DFMEA, RCA, validation, calibration.",
    },
    "03_Daimler_Mercedes_Commercial_EV_SDV": {
        "resume_label": "Commercial_EV_Software_Electronics_SDV",
        "headline": "Commercial EV Software & Electronics Lead | Embedded SDV Systems",
        "target_roles": [
            "Embedded Software Architect / Engineer - Commercial Vehicles",
            "EV Powertrain Software Lead - Trucks & Buses",
            "Software and Electronics Engineer",
            "Vehicle Systems Integration Lead - Electrification",
        ],
        "summary": "Commercial EV and hybrid powertrain software leader with 8 years delivering truck, bus, and retrofit electrification systems from architecture through validation and ICAT/ARAI homologation. Strong fit for engineering roles requiring embedded automotive software, MATLAB/Simulink, CAN/LIN/Ethernet mindset, AUTOSAR basics, ISO 26262, software-defined vehicle readiness, supplier coordination, and global team collaboration. Directly supported 5T EV architecture and certified commercial EV and P4 hybrid programmes.",
        "strengths": [
            "Strong commercial vehicle fit: LCV/LPT EV, 5T EV, 16T P4 hybrid bus, and 6T GVW EV ARAI certification exposure across trucks and buses.",
            "Developed eVCU/BMS application software and system logic for power-up/power-down sequencing, vehicle operating states, DCDC/OBC/PDU integration, charging controller interfaces, diagnostics, and HV-LV handshakes.",
            "Comfortable with software/electronics requirements: MATLAB, CAN, embedded software, functional safety, AUTOSAR basics, system decomposition, global coordination, supplier quality, and change-ready product development.",
        ],
        "experience_bullets": [
            "Led EV and P4 hybrid powertrain architecture, ECU software, battery engineering, supplier development, validation, and homologation readiness for commercial vehicle platforms.",
            "Supported 5T EV conversion architecture with 80 kW powertrain, 56 kWh / 310V NMC battery system, power electronics interfaces, HV battery integration, vehicle validation, and supplier coordination.",
            "Developed MATLAB/Simulink/Stateflow eVCU and BMS logic for functional requirements, safety requirements, V-model verification artefacts, and ISO 26262 ASIL-aligned control behaviour.",
            "Integrated vehicle communication and software interfaces across CAN-connected battery, charger, DCDC, OBC, motor controller, PDU, telematics, instrument cluster, EPS, BCS, TCS, and HVAC systems.",
            "Managed cross-functional teams and global suppliers while applying DFMEA, RCA, Six Sigma, and validation planning to close integration issues and improve release readiness.",
        ],
        "project_bullets": [
            "5T EV conversion: supported EV architecture, 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, power electronics interfaces, and vehicle validation.",
            "LCV/LPT EV and 6T GVW EV ARAI programme: delivered commercial EV software/system integration exposure directly aligned with electric truck and bus development.",
            "16T P4 hybrid ICAT programme: collaborated with senior OEM/Tier-1 leadership on hybrid bus certification and state government pilot fleet trials.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, Embedded C/C++, CAN, DBC, Vector CANoe, Peak CAN, AUTOSAR basics, ISO 26262, V-model, Git awareness, eVCU, BMS, DCDC, OBC, PDU, HV battery, commercial EV, SDV, trucks, buses, homologation.",
    },
    "04_BMW_Embedded_ECU_Software": {
        "resume_label": "Embedded_ECU_Vehicle_Controls_MBD",
        "headline": "Embedded ECU Software Lead | Vehicle Controls & Model-Based Development",
        "target_roles": [
            "Embedded Software Developer - ECU Platform",
            "Automotive Software Engineer - Vehicle Controls",
            "Software Integration / Systems Engineer - Electronic Control Units",
            "Model-Based Development Engineer",
        ],
        "summary": "Automotive embedded software and EV systems leader targeting ECU software, vehicle controls, software integration, and model-based development roles. Offers 8 years of hands-on eVCU/BMS application software, MATLAB/Simulink/Stateflow controls, embedded C/C++, CAN diagnostics, HV battery integration, and safety-critical EV powertrain delivery. Brings proven system-level experience from commercial EV conversions, P4 hybrid buses, certified HV battery packs, and vehicle electronics including instrument clusters, dual-CAN telematics, and LV/HV PDU development.",
        "strengths": [
            "ECU software fit: practical experience with electronic control unit software, related embedded tools, system decomposition, and vehicle-level software integration.",
            "Developed model-based vehicle control and BMS software across drive readiness, charging, DCDC/OBC/PDU integration, fault handling, diagnostics, and safe operating states.",
            "Delivered embedded vehicle electronics including STM32F4/F1 instrument cluster, Quectel dual-CAN telematics, ultra-capacitor monitoring, and LV/HV power distribution control.",
        ],
        "experience_bullets": [
            "Developed BMS and eVCU application software with MATLAB/Simulink/Stateflow MBD methodology, aligning functional requirements, safety requirements, and verification outputs with V-model discipline.",
            "Implemented vehicle control logic, power-up/power-down sequencing, DCDC/OBC/PDU coordination, charging controller interfaces, HV-LV handshake protocols, and diagnostics for EV and hybrid platforms.",
            "Built component-level and vehicle-level range estimation models and integrated battery, motor, gearbox, controller, EPS, BCS, TCS, and HVAC systems through structured interface definition.",
            "Developed STM32-based instrument cluster, dual-CAN telematics, LV/HV PDU, and ultra-capacitor monitoring/control system, combining embedded software with vehicle electronics requirements.",
            "Applied ISO 26262 awareness, IATF 16949 process thinking, DFMEA, RCA, and Six Sigma methods to improve design robustness and validation efficiency.",
        ],
        "project_bullets": [
            "LCV/LPT EV conversion: ARAI-homologated commercial EV with 80 kW drive and 53 kWh LFP battery; contributed ECU software, eVCU/BMS logic, charging, diagnostics, DCDC/OBC/PDU coordination, and vehicle control.",
            "In-house embedded vehicle electronics: built STM32F4/F1 instrument cluster, Quectel dual-CAN telematics, LV/HV PDU, and ultra-capacitor cell monitoring/control system.",
            "HV battery systems: designed and packaged 5 kWh to 300 kWh battery systems across 72V-800V LFP, NMC, LTO, and ultra-capacitor architectures.",
        ],
        "tools": "Embedded C/C++, MATLAB, Simulink, Stateflow, ECU software, eVCU, BMS, VCU, STM32, NXP, Infineon, CAN, DBC, Vector CANoe, Peak CAN, diagnostics, ISO 26262, AUTOSAR basics, V-model, EV powertrain, HV battery.",
    },
    "05_JLR_LandRover_EDU_EV_Validation": {
        "resume_label": "EV_Powertrain_EDU_Validation",
        "headline": "EV Powertrain Validation & Controls Lead | EDU, Battery & Software",
        "target_roles": [
            "EV Powertrain Validation Engineer",
            "Electric Drive Unit / EDU Validation Engineer",
            "Vehicle Controls & Systems Integration Engineer",
            "Battery / Powertrain Software Validation Lead",
        ],
        "summary": "EV powertrain software and validation leader targeting Electric Drive Unit validation, powertrain testing, vehicle controls, battery integration, and software-enabled test/data workflow roles. Brings 8 years in EV and hybrid commercial platforms with MATLAB/Simulink/Stateflow controls, BMS/eVCU software, HV battery architecture, DCDC/OBC/PDU integration, CAN diagnostics, test planning, DFMEA/RCA, Six Sigma, and ICAT/ARAI homologation. Strong fit for roles requiring powertrain testing at component, subsystem, and system levels plus cross-functional issue resolution.",
        "strengths": [
            "Powertrain validation fit: delivered EV and hybrid platforms from architecture to vehicle validation, certification sign-off support, and issue closure using structured problem-solving tools.",
            "MATLAB/Simulink and data workflow fit: built range estimation models, control strategies, and verification artefacts; comfortable using model/data analysis to support efficient validation decisions.",
            "EDU/electrification fit: integrated motor, gearbox, HV battery, DCDC, OBC, PDU, charger, BMS, controller, and vehicle controls across commercial EV platforms.",
        ],
        "experience_bullets": [
            "Led EV and hybrid system integration across battery, motor, gearbox, DCDC, OBC, PDU, charger, controllers, sensors, telematics, and vehicle control interfaces.",
            "Developed eVCU/BMS application software and operating-state logic using MATLAB/Simulink/Stateflow for drive readiness, charging mode, fault handling, power sequencing, and safe shutdown behaviour.",
            "Supported vehicle-level validation, test planning, fault investigation, CAN diagnostics, calibration behaviour review, supplier issue closure, and homologation evidence preparation.",
            "Applied DFMEA, RCA, DMAIC, Six Sigma, and quality tools to resolve integration issues, reduce validation cycles, and improve design robustness.",
            "Managed EV programme deliverables across internal engineering teams and external suppliers, communicating technical risk and readiness at system, subsystem, and component levels.",
        ],
        "project_bullets": [
            "LCV/LPT EV conversion: ARAI-homologated 6.5T EV with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed ECU logic, charging interfaces, HV battery coordination, diagnostics, and validation support.",
            "5T commercial EV: supported 80 kW EV conversion with 56 kWh / 310V NMC pack, power electronics interfaces, vehicle validation, and supplier coordination.",
            "P4 hybrid bus: integrated super-capacitor electrification kit for a 16T platform, supported ICAT validation, and contributed to 25% fuel-efficiency improvement evidence.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, powertrain validation, Electric Drive Unit, EDU, HV battery, BMS, eVCU, CAN, Vector CANoe, Peak CAN, test planning, instrumentation, DFMEA, RCA, Six Sigma, vehicle integration, ISO 26262, ICAT, ARAI.",
    },
    "06_AVL_EMobility_Simulation_Test_Systems": {
        "resume_label": "EMobility_Simulation_Test_Systems",
        "headline": "E-Mobility, Simulation & Test Systems Lead | EV Software Engineering",
        "target_roles": [
            "E-Mobility Systems Engineer",
            "Software Development Engineer - EV Test / Simulation",
            "Battery / E-Drive Test Systems Engineer",
            "Model-Based Development / System Integration Engineer",
        ],
        "summary": "E-mobility software and systems leader targeting roles in electrification, simulation, test/validation systems, charging solutions, and vehicle efficiency engineering. Offers 8 years delivering EV/hybrid systems with MATLAB/Simulink/Stateflow MBD, range estimation models, BMS/eVCU software, charging controller logic, HV battery systems, DCDC/OBC/PDU integration, and validation/homologation support. Strongly aligned with engineering, simulation, testing, electrification, software, AI/automation, battery/e-drive/inverter validation, and model-based virtual development before hardware build.",
        "strengths": [
            "Built MBD models for component-level and vehicle-level EV/hybrid range estimation, supporting virtual design decisions, vehicle efficiency studies, and validation planning.",
            "Led EV charging controller logic and system integration across charger, BMS, eVCU/VCU, DCDC, OBC, PDU, HV battery, contactors, and vehicle safety states.",
            "Designed HV battery systems and certified battery packs across LFP, NMC, LTO, and ultra-capacitor chemistries for work around batteries, e-motors, inverters, simulation, testing, and charging solutions.",
        ],
        "experience_bullets": [
            "Developed MATLAB/Simulink/Stateflow control models for eVCU/BMS functions, range estimation, charging logic, and EV/hybrid operating-state behaviour.",
            "Integrated EV components and test/validation workflows across battery, motor, gearbox, DCDC, OBC, PDU, charger, controller, telematics, and thermal systems.",
            "Applied DFMEA, RCA, Six Sigma, and validation evidence reviews to improve design robustness and shorten test/debug cycles across EV programmes.",
            "Led supplier coordination and hardware-software interface definition for battery, motor, gearbox, DCDC, EPS, BCS, TCS, HVAC, and controller systems.",
            "Supported vehicle-level validation, certification documentation, and homologation readiness for ICAT/ARAI-certified EV and hybrid programmes.",
        ],
        "project_bullets": [
            "E-mobility models: created EV/hybrid range estimation and control models for component and vehicle-level scenarios using MATLAB/Simulink.",
            "Battery and charging systems: delivered 5 kWh to 300 kWh HV battery systems, AC/DC charging controller logic, DCDC/OBC/PDU integration, and BTMS/HVAC validation.",
            "Commercial EV programmes: supported LCV/LPT EV, 5T EV, 16T P4 hybrid bus, sedan, and light commercial EV conversion programmes through integration and validation.",
        ],
        "tools": "MATLAB, Simulink, Stateflow, model-based simulation, range estimation, e-mobility, EV test systems, BMS, eVCU, charging systems, DCDC, OBC, PDU, battery testing, e-drive, inverter, CAN, diagnostics, DFMEA, RCA, Six Sigma.",
    },
    "07_Bosch_Automotive_Embedded_Systems": {
        "resume_label": "Automotive_Embedded_Systems_Integration",
        "headline": "Automotive Embedded Systems & EV Integration Lead | Mobility Systems",
        "target_roles": [
            "System Integration Engineer - Automotive Applications",
            "Embedded Software Engineer - Automotive Products",
            "EV / Cross-Domain Vehicle Systems Engineer",
            "BMS / ECU Software Integration Lead",
        ],
        "summary": "Automotive embedded systems and EV software leader targeting roles involving cross-domain automotive products, embedded software, hardware electronics integration, future vehicle architectures, and system integration. Brings 8 years of eVCU/BMS software, MATLAB/Simulink/Stateflow MBD, embedded C/C++, STM32/NXP/Infineon exposure, CAN diagnostics, HV battery systems, DCDC/OBC/PDU integration, charging controller logic, and vehicle-level validation. Strong fit for roles connecting hardware electronics, embedded software, software architecture, automotive product development, and global engineering delivery.",
        "strengths": [
            "System integration fit: developed cross-domain EV systems linking hardware electronics, embedded software, CAN communication, power electronics, battery systems, vehicle controls, and validation.",
            "Embedded systems fit: developed STM32-based instrument cluster, dual-CAN telematics, ultra-capacitor monitoring/control, LV/HV PDU, and eVCU/BMS application software.",
            "Automotive product fit: delivered SOR, supplier coordination, validation evidence, quality tools, and ICAT/ARAI homologation readiness for commercial EV and hybrid programmes.",
        ],
        "experience_bullets": [
            "Integrated hardware electronics and embedded software across HV battery, BMS/eVCU, DCDC, OBC, PDU, motor controller, telematics, sensors, and vehicle controls.",
            "Developed MATLAB/Simulink/Stateflow eVCU and BMS application software for functional requirements, safety requirements, diagnostics, fault handling, and safe vehicle operating states.",
            "Built STM32F4/F1 instrument cluster, Quectel dual-CAN telematics, LV/HV PDU, and ultra-capacitor cell monitoring/control system for in-house EV component development.",
            "Defined component/system specifications and SOR documents for supplier sourcing and integration across battery, motor, gearbox, DCDC, EPS, BCS, TCS, HVAC, and controller systems.",
            "Used CAN tools, diagnostics, DFMEA, RCA, Six Sigma, ISO 26262 awareness, and IATF 16949 process thinking to support robust automotive product development.",
        ],
        "project_bullets": [
            "LCV/LPT EV conversion: ARAI-homologated commercial EV with eVCU/BMS logic, vehicle control, charging interfaces, DCDC/OBC/PDU integration, diagnostics, and HV battery coordination.",
            "Embedded component suite: delivered LV/HV PDU, STM32 instrument cluster, dual-CAN telematics, ultra-capacitor monitoring/control, and battery-related software interfaces.",
            "HV battery systems: packaged and integrated 5 kWh to 300 kWh packs across LFP, NMC, LTO, and ultra-capacitor technologies with BTMS/HVAC validation.",
        ],
        "tools": "Embedded C/C++, STM32, Microchip PIC, NXP, Infineon, MATLAB, Simulink, Stateflow, BMS, eVCU, CAN, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, DBC, diagnostics, software architecture, system integration, ISO 26262, IATF 16949.",
    },
    "08_Continental_Software_System_Engineering": {
        "resume_label": "Software_System_Engineering_EV_Controls",
        "headline": "Software & System Engineering Lead | EV Controls, BMS and Vehicle Integration",
        "target_roles": [
            "Software Engineering / System Engineering - Automotive",
            "EV Systems Integration Engineer",
            "Embedded Software Engineer - Mobility Systems",
            "BMS / Vehicle Controls Lead",
        ],
        "summary": "Automotive software and systems engineering leader targeting roles in software engineering, system engineering, safer mobility, sustainable vehicle technologies, and future mobility solutions. Brings 8 years across EV/hybrid architecture, eVCU/BMS software, MATLAB/Simulink/Stateflow MBD, embedded C/C++, CAN diagnostics, HV battery systems, DCDC/OBC/PDU integration, vehicle validation, and quality-led development. Strong fit for roles needing cross-functional systems thinking, embedded software, EV powertrain integration, safety/process discipline, and customer-ready mobility products.",
        "strengths": [
            "Software/system engineering fit: led EV architecture, ECU application software, component specifications, supplier interfaces, validation planning, and homologation readiness.",
            "Mobility technology fit: delivered EV and hybrid systems that improve efficiency, safety, and sustainability across commercial vehicle applications.",
            "Process fit: combined ISO 26262 functional safety awareness, IATF 16949 quality management, DFMEA, RCA, Six Sigma, and validation evidence into structured automotive delivery.",
        ],
        "experience_bullets": [
            "Led EV and P4 hybrid architecture, ECU application software, HV battery engineering, supplier development, validation, and certification readiness for commercial vehicle platforms.",
            "Developed MATLAB/Simulink/Stateflow eVCU/BMS logic for vehicle control, charging, DCDC/OBC/PDU coordination, fault handling, diagnostics, and safe operating states.",
            "Defined system and component specifications for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, EPS, BCS, TCS, HVAC, and embedded vehicle electronics.",
            "Managed cross-functional engineering and supplier teams to improve component readiness, validation efficiency, system integration, and release-quality software/hardware interfaces.",
            "Applied DFMEA, RCA, DMAIC, SPC, and control charts to reduce risk, improve supplier quality, and support automotive compliance deliverables.",
        ],
        "project_bullets": [
            "Commercial EV systems: LCV/LPT EV, 5T EV, sedan, and light commercial EV conversions with HV battery, charger, DCDC, controller, vehicle control, and validation integration.",
            "P4 hybrid bus: 16T ICAT-certified hybrid programme with super-capacitor pack and 25% fuel-efficiency improvement validation.",
            "In-house software/electronics: LV/HV PDU, STM32 instrument cluster, dual-CAN telematics, ultra-capacitor cell monitoring/control, and EV range estimation models.",
        ],
        "tools": "Software Engineering, System Engineering, MATLAB, Simulink, Stateflow, Embedded C/C++, BMS, eVCU, VCU, CAN, DBC, diagnostics, HV battery, DCDC, OBC, PDU, vehicle integration, ISO 26262, IATF 16949, DFMEA, RCA, Six Sigma.",
    },
}


def resume_view(target):
    data = dict(target)
    data.update(RESUME_OVERRIDES[target["folder"]])
    validate_resume_terms(data)
    return data


def validate_resume_terms(target):
    resume_strings = [
        target["resume_label"],
        target["headline"],
        target["summary"],
        target["tools"],
        *target["target_roles"],
        *target["strengths"],
        *target["experience_bullets"],
        *target["project_bullets"],
    ]
    violations = []
    for text in resume_strings:
        lowered = text.lower()
        for term in FORBIDDEN_RESUME_TERMS:
            if term in lowered:
                violations.append((term, text))
    if violations:
        details = "\n".join(f"- {term}: {text}" for term, text in violations)
        raise ValueError(f"Target-company term found in resume content:\n{details}")


def clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def set_font(run, name="Arial", size=9.3, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=2.5, line=1.03):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_rule(paragraph, color="D9D9D9", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=7, after=3, line=1.0)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=True, color="1F4D78")
    add_bottom_rule(p)
    return p


def add_body(doc, text, size=9.1, after=2.5):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.03)
    set_font(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=8.85):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=1.2, line=1.01)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    set_font(p.add_run(text), size=size)
    return p


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.50)
    section.bottom_margin = Inches(0.50)
    section.left_margin = Inches(0.60)
    section.right_margin = Inches(0.60)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    for style_name, size in [("Normal", 9.1), ("List Bullet", 8.85)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)


def add_role_header(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=0, line=1.0)
    set_font(p.add_run("Technical Manager - Product Development"), size=9.8, bold=True)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=2, line=1.0)
    set_font(p2.add_run(f"{BASE_FACTS['company']} | {BASE_FACTS['dates']}"), size=8.7, color="555555")


def build_resume(target, out_dir: Path):
    target = resume_view(target)
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, after=1, line=1.0)
    set_font(title.add_run(NAME.upper()), size=17.2, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=1, line=1.0)
    set_font(subtitle.add_run(target["headline"]), size=9.9, bold=True, color="1F4D78")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=5, line=1.0)
    set_font(contact.add_run(CONTACT), size=8.4, color="555555")

    add_heading(doc, "PROFESSIONAL SUMMARY")
    add_body(doc, target["summary"], size=9.1)

    add_heading(doc, "TARGET ROLE FIT")
    for role in target["target_roles"]:
        add_bullet(doc, role, size=8.8)

    add_heading(doc, "CORE STRENGTHS")
    for strength in target["strengths"]:
        add_bullet(doc, strength)

    add_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_header(doc)
    for bullet in target["experience_bullets"]:
        add_bullet(doc, bullet)

    add_heading(doc, "SELECTED PROGRAMMES & PROOF")
    for bullet in target["project_bullets"]:
        add_bullet(doc, bullet)

    add_heading(doc, "TECHNICAL TOOLKIT & CERTIFICATIONS")
    add_body(doc, target["tools"], size=8.8, after=2)
    for cert in BASE_FACTS["certs"]:
        add_bullet(doc, cert, size=8.8)

    add_heading(doc, "EDUCATION")
    add_body(doc, BASE_FACTS["education"], size=8.9)

    filename_base = f"Kuldeep_Singh_{target['resume_label']}_Resume_R02"
    docx_path = out_dir / f"{filename_base}.docx"
    txt_path = out_dir / f"{filename_base}.txt"
    doc.save(docx_path)

    lines = [
        NAME.upper(),
        target["headline"],
        CONTACT,
        "",
        "PROFESSIONAL SUMMARY",
        target["summary"],
        "",
        "TARGET ROLE FIT",
        *[f"- {x}" for x in target["target_roles"]],
        "",
        "CORE STRENGTHS",
        *[f"- {x}" for x in target["strengths"]],
        "",
        "PROFESSIONAL EXPERIENCE",
        "Technical Manager - Product Development",
        f"{BASE_FACTS['company']} | {BASE_FACTS['dates']}",
        *[f"- {x}" for x in target["experience_bullets"]],
        "",
        "SELECTED PROGRAMMES & PROOF",
        *[f"- {x}" for x in target["project_bullets"]],
        "",
        "TECHNICAL TOOLKIT & CERTIFICATIONS",
        target["tools"],
        *[f"- {x}" for x in BASE_FACTS["certs"]],
        "",
        "EDUCATION",
        BASE_FACTS["education"],
        "",
    ]
    txt_path.write_text("\n".join(lines), encoding="utf-8")
    return docx_path, txt_path


def build_apply_notes(target, out_dir: Path):
    notes = [
        f"# {target['doc_label'].replace('_', ' ')} - Application Notes",
        "",
        "## Best-Fit Roles",
        *[f"- {x}" for x in target["target_roles"]],
        "",
        "## Where To Apply",
        *[f"- [{label}]({url})" for label, url in target["apply_links"]],
        "",
        "## Resume Positioning",
        *[f"- {x}" for x in target["research_notes"]],
        "",
        "## Portal Search Terms",
        target["search_terms"],
        "",
        "## How To Use This Folder",
        "- Upload the `.docx` file when the portal accepts Word resumes.",
        "- Upload the `.pdf` file when the portal asks for a fixed-format resume.",
        "- Use the `.txt` file to paste into ATS text boxes or recruiter forms.",
        "- Keep the filename with `R02` so version history stays clean.",
        "",
        "## Research Sources",
        *[f"- {url}" for url in target["sources"]],
        "",
    ]
    (out_dir / "WHERE_TO_APPLY.md").write_text("\n".join(notes), encoding="utf-8")


def build_matrix(out_root: Path):
    lines = [
        "# R02 Targeted Resume Pack",
        "",
        "Each target folder contains a tailored `.docx` resume, ATS/plain-text copy, rendered PDF, and `WHERE_TO_APPLY.md` with application links and targeting notes.",
        "",
        "Resume files are role-targeted and company-neutral. Target-company names stay in folder names, application notes, and research links, not inside the resume content or resume filename.",
        "",
        "ATS terms are woven into the resume summary, strengths, experience, programmes, toolkit, and certifications rather than placed in a standalone keyword section.",
        "",
        "| Target Folder | Primary Target | Best Role Families |",
        "|---|---|---|",
    ]
    for t in TARGETS:
        lines.append(
            f"| `{t['folder']}` | {t['doc_label'].replace('_', ' ')} | "
            + "; ".join(t["target_roles"][:2])
            + " |"
        )
    lines.extend(
        [
            "",
            "## Version",
            "",
            "- Resume package version: `R02`",
            "- Generated from Kuldeep Singh executive resume and current public target-company career research.",
            "- Keep all future updates as `R03`, `R04`, etc. without overwriting previous versions.",
            "",
        ]
    )
    (out_root / "README.md").write_text("\n".join(lines), encoding="utf-8")


def main():
    clean_dir(OUTPUT_ROOT)
    clean_dir(REPO_R02)
    generated = []
    for target in TARGETS:
        out_dir = OUTPUT_ROOT / target["folder"]
        repo_dir = REPO_R02 / target["folder"]
        out_dir.mkdir(parents=True, exist_ok=True)
        repo_dir.mkdir(parents=True, exist_ok=True)

        docx_path, txt_path = build_resume(target, out_dir)
        build_apply_notes(target, out_dir)
        for file in out_dir.iterdir():
            if file.is_file():
                shutil.copy2(file, repo_dir / file.name)
        generated.append(docx_path)

    build_matrix(OUTPUT_ROOT)
    shutil.copy2(OUTPUT_ROOT / "README.md", REPO_R02 / "README.md")
    print(OUTPUT_ROOT)
    for docx_path in generated:
        print(docx_path)


if __name__ == "__main__":
    main()
