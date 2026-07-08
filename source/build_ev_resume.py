from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kuldeepsingh/Documents/Codex/2026-07-08/en")
OUTPUTS = ROOT / "outputs"
OUTPUTS.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUTPUTS / "Kuldeep_Singh_EV_ECU_Software_Lead_ATS_Resume.docx"
TXT_PATH = OUTPUTS / "Kuldeep_Singh_EV_ECU_Software_Lead_ATS_Resume.txt"

NAME = "Kuldeep Singh"
HEADLINE = "Technical Manager | MATLAB/Simulink EV ECU Software & Electrification Lead"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"


SECTIONS = [
    (
        "PROFESSIONAL SUMMARY",
        [
            "Senior electrification and EV ECU software leader with 8 years of experience delivering safety-critical EV and hybrid powertrain systems from concept and SOR definition through ICAT/ARAI homologation and production readiness. Specialized in MATLAB/Simulink/Stateflow model-based development for eVCU and BMS application software, vehicle control, charging control, DCDC/OBC/PDU integration, power-up/power-down sequencing, HV-LV handshake logic, and V-model verification under ISO 26262 ASIL guidelines. Led complete EV retrofit and commercial vehicle electrification programs for TATA, Mahindra, Daimler Fuso, and government fleet platforms, including ARAI-homologated TATA 407 LCV/LPT EV work.",
        ],
    ),
    (
        "CORE TECHNICAL SKILLS",
        [
            "EV ECU Software: eVCU/VCU application software, BMS application software, charging controller, AC/DC charging architectures, DCDC converter integration, OBC integration, PDU logic, contactor/precharge sequencing, drive enable logic, power-up/power-down sequencing, vehicle control logic, fault handling, diagnostics, inter-system handshake protocols.",
            "Model-Based Development: MATLAB, Simulink, Stateflow, MBD, requirements-to-model workflows, V-model software lifecycle, functional requirements, safety requirements, range estimation modelling, MIL/SIL/HIL readiness, Simulink Coder/Embedded Coder concepts, reusable subsystem design.",
            "Embedded & Vehicle Networks: Embedded C/C++, STM32 8/16/32-bit, Microchip PIC, NXP, Infineon, CAN bus, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, DBC, diagnostics, calibration, ECU integration, telematics, instrument cluster development.",
            "EV Powertrain & Battery: HV battery architecture, 72V-800V systems, LFP, NMC, LTO, ultra-capacitor, BMU/BMS, BTMS/HVAC, motor, gearbox, DCDC, OBC, PDU, EPS, BCS, TCS, thermal management, commercial EV conversion, P4 hybrid systems.",
            "Compliance & Leadership: ISO 26262, IATF 16949, DFMEA, RCA, Six Sigma Black Belt, DMAIC, SPC, control charts, supplier development, SOR documentation, cross-functional team leadership, ICAT/ARAI homologation, programme delivery.",
        ],
    ),
    (
        "PROFESSIONAL EXPERIENCE",
        [
            ("role", "Technical Manager - Product Development", "IX Energy Pvt. Ltd., Noida, India | Jul 2018 - Present"),
            "Lead end-to-end EV and P4 hybrid powertrain architecture, ECU application software, battery engineering, supplier management, vehicle integration, validation, and homologation readiness for commercial vehicle electrification programs.",
            "Architected and delivered complete EV and P4 hybrid systems from SOR definition through ICAT/ARAI homologation for TATA, Mahindra, and Daimler Fuso platforms.",
            "Developed BMS and eVCU application software using MATLAB/Simulink/Stateflow MBD methodology, covering functional requirements, safety requirements, V-model verification artefacts, and ISO 26262 ASIL-aligned logic.",
            "Implemented power-up/power-down sequencing, vehicle control logic, DCDC/OBC/PDU integration logic, charging controller interfaces, HV-LV handshake protocols, fault handling, and inter-system operating-state logic.",
            "Developed EV charging controller logic as technical lead for AC and DC charging architectures, aligning software behaviour with charger, BMS, VCU/eVCU, contactor, and vehicle safety requirements.",
            "Defined component and system specifications for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, EPS, BCS, TCS, and HVAC; authored hardware and software SOR documents for sourcing and integration.",
            "Built MBD models for component-level and vehicle-level EV/hybrid range estimation; applied DFMEA, RCA, and Six Sigma Black Belt methods to improve design robustness and validation efficiency.",
            "Managed cross-functional teams and global supplier networks across battery, motor, gearbox, DCDC, EPS, BCS, TCS, and HVAC to maintain component readiness against programme milestones.",
            "Collaborated directly with senior technical leadership at TATA AutoComp Systems on 16T GVW P4 Hybrid ICAT certification and Bharat Forge/Kalyani Group on 6T GVW EV ARAI certification.",
        ],
    ),
    (
        "SELECTED EV / ECU PROGRAMMES",
        [
            "TATA 407 LCV / LPT EV Conversion - ARAI homologated 6.5T commercial EV with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed MATLAB/Simulink eVCU/BMS logic, vehicle control, charging interfaces, DCDC/PDU/OBC integration, diagnostics, safety logic, and homologation-readiness support.",
            "TATA 1512 P4 Hybrid Bus Programme - ICAT certified 16T GVW hybrid bus with super-capacitor pack; developed and integrated P4 hybrid electrification kit and supported validation showing 25% fuel-efficiency improvement.",
            "Daimler Fuso 5T EV Platform - supported EV conversion architecture, 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, HV battery integration, power electronics interfaces, vehicle validation, and supplier coordination.",
            "Ambassador Car & Mahindra Supro EV Conversions - delivered EV conversion engineering for 30 kW powertrain, 15 kWh / 72V LFP battery, 100 km range, charging, vehicle controls, validation, and integration requirements.",
            "In-House EV Component Development - developed LV/HV PDU, STM32F4/F1-based instrument cluster, Quectel dual-CAN telematics unit, and ultra-capacitor cell monitoring and control system.",
        ],
    ),
    (
        "HV BATTERY, BMS & POWER ELECTRONICS",
        [
            "Designed and packaged HV battery systems from 5 kWh to 300 kWh using LFP, NMC, LTO, and ultra-capacitor chemistries across 72V-800V single-string and multi-string architectures.",
            "Led BMU development on Toshiba automotive module and BTMS/HVAC design, integration, and validation for 3.3-5 kW thermal management systems.",
            "Delivered certified battery packs including 53 kWh/332V LFP electric truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC MP-Co pack, and 0.5 kWh/400V ultra-capacitor hybrid bus pack.",
            "Integrated battery, DCDC, OBC, PDU, controller, motor, gearbox, EPS, BCS, TCS, and HVAC systems through structured hardware-software interface definition and supplier coordination.",
        ],
    ),
    (
        "EDUCATION & CERTIFICATIONS",
        [
            "B.Tech - Mechanical Engineering, VIT University, Vellore, Tamil Nadu | 2014 - 2018 | First Class, 75%",
            "Special Achiever Award, 2017 & 2018 - University-level recognition for outstanding contribution to INR 50 lakh funded research projects.",
            "Six Sigma Black Belt - Certified",
            "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
            "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
            "AWS IoT & SIMULINK - Amazon Web Services / MathWorks Certified",
        ],
    ),
    (
        "TARGET ATS KEYWORDS",
        [
            "MATLAB, Simulink, Stateflow, Model-Based Design, MBD, eVCU, VCU, ECU Application Software, BMS Software, EV Battery, HV Battery, DCDC Converter, OBC, On-Board Charger, PDU, Charging Controller, AC Charging, DC Charging, CCS2, Vehicle Control, Power-Up Sequence, Power-Down Sequence, Contactor Control, Precharge, Diagnostics, CAN, CANoe, Peak CAN, Bus Master, DBC, Embedded C, Embedded C++, STM32, NXP, Infineon, ISO 26262, ASIL, IATF 16949, AUTOSAR Basics, V-Model, MIL, SIL, HIL, DFMEA, RCA, Six Sigma Black Belt, ARAI, ICAT, Homologation, EV Retrofit, Commercial EV, P4 Hybrid, TATA 407, TATA 407 LPT, TATA Motors, TATA AutoComp, VinFast, Tesla, Mahindra, Daimler Fuso, Bharat Forge, Kalyani Group.",
        ],
    ),
]


def set_font(run, name="Arial", size=9.5, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=3, line=1.05):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_rule(paragraph, color="D9D9D9", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_heading(doc, text):
    p = doc.add_paragraph()
    set_spacing(p, before=7, after=3, line=1.0)
    r = p.add_run(text)
    set_font(r, size=10.5, bold=True, color="1F4D78")
    add_bottom_rule(p)
    return p


def add_body(doc, text, after=3, size=9.2):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.04)
    r = p.add_run(text)
    set_font(r, size=size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=1.5, line=1.02)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.14)
    r = p.add_run(text)
    set_font(r, size=9.0)
    return p


def add_role(doc, title, meta):
    p = doc.add_paragraph()
    set_spacing(p, before=3, after=0, line=1.0)
    r = p.add_run(title)
    set_font(r, size=9.8, bold=True)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=2, line=1.0)
    r2 = p2.add_run(meta)
    set_font(r2, size=8.8, color="555555")


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.3)
    section.footer_distance = Inches(0.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    normal.font.size = Pt(9.2)

    bullet = doc.styles["List Bullet"]
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
    bullet.font.size = Pt(9.0)


def build_docx():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, after=1, line=1.0)
    set_font(title.add_run(NAME.upper()), size=17.5, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=1, line=1.0)
    set_font(subtitle.add_run(HEADLINE), size=10.2, bold=True, color="1F4D78")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=5, line=1.0)
    set_font(contact.add_run(CONTACT), size=8.5, color="555555")

    for heading, items in SECTIONS:
        if heading == "HV BATTERY, BMS & POWER ELECTRONICS":
            doc.add_page_break()
        add_heading(doc, heading)
        for item in items:
            if isinstance(item, tuple) and item and item[0] == "role":
                add_role(doc, item[1], item[2])
            elif heading in {"PROFESSIONAL SUMMARY", "TARGET ATS KEYWORDS"}:
                add_body(doc, item, after=2, size=9.0 if heading == "TARGET ATS KEYWORDS" else 9.2)
            elif heading == "CORE TECHNICAL SKILLS":
                label, _, rest = item.partition(":")
                p = doc.add_paragraph()
                set_spacing(p, after=2, line=1.02)
                r1 = p.add_run(label + ": ")
                set_font(r1, size=8.9, bold=True)
                r2 = p.add_run(rest.strip())
                set_font(r2, size=8.9)
            elif heading == "EDUCATION & CERTIFICATIONS":
                add_bullet(doc, item)
            else:
                add_bullet(doc, item)

    doc.save(DOCX_PATH)


def build_text():
    lines = [NAME.upper(), HEADLINE, CONTACT, ""]
    for heading, items in SECTIONS:
        lines.append(heading)
        for item in items:
            if isinstance(item, tuple) and item and item[0] == "role":
                lines.append(item[1])
                lines.append(item[2])
            elif heading in {"PROFESSIONAL SUMMARY", "TARGET ATS KEYWORDS"}:
                lines.append(item)
            else:
                lines.append(f"- {item}")
        lines.append("")
    TXT_PATH.write_text("\n".join(lines).strip() + "\n", encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_text()
    print(DOCX_PATH)
    print(TXT_PATH)
