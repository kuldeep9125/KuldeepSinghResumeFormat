from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kuldeepsingh/Documents/Codex/2026-07-08/en")
OUTPUT_ROOT = ROOT / "outputs" / "R05_Project_Verified_Full_Resume"
REPO_ROOT = ROOT / "work" / "KuldeepSinghResumeFormat"
REPO_R05 = REPO_ROOT / "resumes" / "R05"
LOCAL_ROOT = Path("/Users/kuldeepsingh/Downloads/RESUME/R05_Project_Verified_Full_Resume")

VERSION = "R05"
FILE_BASE = "Kuldeep_Singh_Project_Verified_EV_Software_Technical_Lead_Resume_R05"

NAME = "Kuldeep Singh"
HEADLINE = "Automotive EV Software & Technical Lead - Simulink/Stateflow, eVCU/BMS, CAN, HV Battery"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"
COMPANY = "IX Energy Pvt. Ltd., Noida, India"
ROLE = "Technical Manager - Product Development"
DATES = "Jul 2018 - Present"
EDUCATION = "B.Tech - Mechanical Engineering, VIT University, Vellore, Tamil Nadu | 2014 - 2018 | First Class, 75%"

SUMMARY = (
    "Automotive EV software and technical lead with 8 years delivering commercial EV, hybrid bus, HV battery, "
    "embedded electronics, and vehicle integration programmes from requirements and architecture to model-based "
    "software, generated build outputs, validation evidence, and ICAT/ARAI homologation. Hands-on manager across "
    "MATLAB/Simulink/Stateflow, eVCU/BMS logic, CAN/J1939 diagnostics, HV battery systems, charging, DCDC/OBC/PDU, "
    "supplier readiness, requirements traceability, V-model delivery, functional-safety discipline, and cross-functional release."
)

LEAD_ROLE_FIT = [
    "Lead roles: EV Software Lead, eVCU/BMS Lead, Vehicle Controls Lead, Technical Manager, Systems Integration Lead, MBD/Controls Lead.",
    "Domain proof: certified commercial EVs, P4 hybrid bus, NIDEC G02 + Microvast hybrid control model, 5 kWh to 300 kWh HV battery systems, AC/DC charging, embedded electronics, telematics, and vehicle validation.",
    "Delivery strength: converts customer/supplier requirements into Simulink/Stateflow logic, CAN databases, calibration outputs, bench/vehicle tests, issue closure, and release-ready evidence.",
]

PROJECT_VERIFIED_PROOF = [
    "Independently developed and iterated the NIDEC G02 + Microvast hybrid control programme in MATLAB/Simulink/Stateflow, with final model revisions, harness metadata, and generated target artifacts verified from the provided project archive.",
    "Implemented hybrid control logic covering torque/speed requests, assist/regen strategies, temperature derate updates, fault handling, drive readiness, BMS state inputs, SOC, charging limits, precharge, contactor, and safe operating behaviour.",
    "Built CAN/J1939 communication layers using DBC/ECOCAN-style message definitions for VCU, BMS, motor controller, charger, telematics, remote access, diagnostics, battery status, motor status, and vehicle status signals.",
    "Produced calibration/build evidence including MOT firmware outputs, A2L measurement/calibration files, MF4 test logs, code-log sheets, and model revisions from R01 through R10 final for bench and vehicle validation.",
    "Integrated NIDEC motor controller and Microvast battery interfaces with BMS charge logic, torque/speed control, telematics, DTC/DM01-style diagnostic signals, fault reset, MCU enable, active discharge, precharge, and temperature behaviour.",
]

EXPERIENCE = [
    "Lead EV and P4 hybrid product development across ECU application software, HV battery engineering, power electronics, supplier development, vehicle integration, validation, and homologation readiness.",
    "Develop BMS/eVCU/VCU application software in MATLAB/Simulink/Stateflow, translating functional requirements and safety constraints into drive, charge, assist, regen, diagnostic, fault, derate, and safe-state logic.",
    "Define SORs, component specifications, CAN signal interfaces, test evidence, and supplier deliverables for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, EPS, BCS, TCS, HVAC, telematics, instrumentation, and embedded electronics.",
    "Architect certified commercial EV systems from requirement capture through ICAT/ARAI outputs, linking hardware, software, supplier readiness, validation evidence, vehicle acceptance, and release documentation.",
    "Design/package 5 kWh to 300 kWh HV battery systems across 72V-800V LFP, NMC, LTO, and ultra-capacitor chemistries, including BMU/BMS logic, contactor/precharge, BTMS/HVAC, charging handshake, and pack integration.",
    "Use CAN diagnostics, DBC review, MF4/test-data analysis, bench/vehicle testing, calibration behaviour review, DFMEA/RCA, and MIL/SIL/HIL-ready model practices to reduce validation cycles and close issues.",
    "Manage internal engineering teams and external suppliers across mechanical, electrical, embedded software, battery, power electronics, quality, and vehicle integration workstreams; communicate risk, readiness, and closure to leadership.",
]

RELEASE_VALIDATION_EVIDENCE = [
    "Maintained revision discipline across Simulink model variants, code-log updates, supplier interface changes, generated outputs, and validation builds so engineering changes were traceable from issue to release evidence.",
    "Used MF4 logs, CAN signal review, DBC checks, bench testing, and vehicle trials to tune assist/regen calibration, temperature derate behaviour, charging limits, fault reset, MCU enable, and precharge/active-discharge readiness.",
    "Connected model-level logic to vehicle-level outcomes: drive enable, gear state, torque limits, speed limits, BMS mode, pack current/voltage, charge stop, SOC, LV voltage, DTE, telematics status, and fault reporting.",
    "Balanced hands-on implementation with supplier coordination across NIDEC motor controller, Microvast battery/BMS, charger, telematics, power electronics, and vehicle integration stakeholders.",
]

COMMERCIAL_OUTCOMES = [
    "Delivered engineering work tied to certified road programmes, including ARAI/ICAT evidence, commercial EV integration, hybrid bus certification readiness, and vehicle-level validation closure.",
    "Supported quantified outcomes: 119 km ARAI commercial EV range evidence, 140 km 5T EV target support, and 25% fuel-efficiency improvement validation on the P4 hybrid bus platform.",
    "Scaled technical ownership from model-level logic to full vehicle systems: HV battery, motor controller, charger, DCDC/OBC/PDU, telematics, instrumentation, diagnostics, supplier readiness, and release documentation.",
]

PROGRAMMES = [
    "NIDEC G02 + Microvast hybrid programme: Simulink/Stateflow hybrid control model with torque/speed, assist/regen, BMS charge, CAN/J1939, telematics, diagnostics, MF4 logs, MOT, and A2L build/calibration evidence.",
    "16T P4 hybrid bus: ICAT-certified platform with super-capacitor energy storage and validated 25% fuel-efficiency improvement; supported system integration, pilot fleet trials, validation evidence, and certification readiness.",
    "6.5T LCV/LPT commercial EV: ARAI-homologated platform with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed eVCU/BMS logic, DCDC/PDU/OBC integration, diagnostics, safety logic, and validation.",
    "5T commercial EV: supported 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, HV battery integration, power electronics interfaces, supplier coordination, and vehicle validation.",
    "Multi-platform EV conversions: supported sedan and light commercial EV conversions across HV battery, charger, DCDC, controller, vehicle controls, diagnostics, CAN integration, and validation.",
]

PORTFOLIO = [
    "Battery and charging: 53 kWh/332V LFP truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC pack, 0.5 kWh/400V ultra-capacitor pack, AC/DC charging controller logic, charger/BMS/VCU coordination, HV-LV handshake, contactor/precharge, and fault response.",
    "Control and diagnostics: torque request, speed limit, gear, brake, accelerator, MCU enable, fault reset, active discharge, FailGrade, motor speed/torque/current, DC voltage/current, DTC/SPN-style diagnostics, remote access, and vehicle status interfaces.",
    "Embedded electronics: LV/HV PDU, STM32F4/F1 instrument cluster, Quectel dual-CAN telematics unit, ultra-capacitor cell monitoring/control system, battery-related software interfaces, and CAN-based diagnostics.",
    "Validation evidence: model revisions, Simulink harness, Stateflow logic, DBC/CAN message review, MF4 logs, code-log updates, temperature derate tuning, assist/regen calibration, bench/vehicle testing, and release-quality documentation.",
]

TECHNICAL_TOOLKIT = [
    "Software and MBD: MATLAB, Simulink, Stateflow, Model-Based Design, Embedded C/C++, generated build artifacts, calibration/measurement files, requirements traceability, V-model, MIL/SIL/HIL readiness.",
    "Automotive software: eVCU, VCU, BMS, vehicle controls, operating-state logic, diagnostics, fault handling, derating, safe-state behaviour, calibration, validation, release evidence, software/hardware integration.",
    "Networks and tools: CAN, J1939, DBC, ECOCAN-style scripts, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, MF4 logs, A2L, MOT, UDS-aware diagnostics, DTC/SPN concepts.",
    "EV systems: HV battery, LFP, NMC, LTO, ultra-capacitor, DCDC, OBC, PDU, charger, contactor, precharge, BTMS/HVAC, motor controller, e-drive/EDU, commercial EV, hybrid powertrain.",
    "Leadership and process: technical manager, supplier development, cross-functional leadership, DFMEA, RCA, DMAIC, SPC, control charts, SOR, ISO 26262, IATF 16949, Six Sigma, AUTOSAR basics, SDV awareness.",
]

IMPLEMENTATION_DEPTH = [
    "Archive scan confirmed 2 final SLX models, 6 MATLAB communication/control files, 1 DBC with 26 messages and 211 signals, 2 code-log sheets, MF4 logs, and generated target-output evidence.",
    "Communication scripts covered VCVCCU J1939, torque/speed control, BMS charging, Microvast battery messages, Zettajoule LTO battery frames, and telematics vehicle/battery/motor status.",
    "DBC and model evidence covered demanded torque/speed, torque limits, motor speed/torque/current, DC voltage/current, FailGrade, MCU enable, active discharge, BMS state/mode, precharge ready, charge stop, and EV fault signals.",
    "Code-log evidence showed assist/regen calibration, temperature-derate updates around 40-42 degC, cut-off voltage changes, and validation-oriented tuning notes rather than only paper-level responsibility.",
    "Presents implementation details at a resume-safe level while preserving strong evidence of model ownership, build artifacts, CAN/database work, calibration updates, and validation data.",
]

LEADERSHIP_SIGNAL = [
    "Lead end-to-end technical delivery across requirements, architecture, model logic, supplier interfaces, calibration, generated artifacts, integration, validation, homologation, and leadership reporting.",
    "Strong match for automotive teams needing a manager who can still read model logic, CAN databases, fault interfaces, battery limits, charger handshakes, and validation data personally.",
    "Comfortable with OEM/Tier-1 language: V-model, requirements traceability, release evidence, DFMEA/RCA, supplier issue closure, ISO 26262 awareness, IATF 16949 discipline, and production-readiness gates.",
]

APPLICATION_SCOPE = [
    "Best-fit applications: Automotive EV Software Lead, Vehicle Controls Lead, eVCU/BMS Lead, Technical Manager, Battery Systems Lead, Systems Integration Lead, and MBD/Validation Lead.",
    "Can contribute at both levels: management of teams/suppliers and direct technical review of Simulink logic, CAN databases, calibration files, logs, diagnostic behaviour, and validation evidence.",
    "Useful for OEM, Tier-1, EV startup, commercial vehicle, battery, power-electronics, and engineering-services roles where software, hardware, validation, and delivery meet.",
]

CERTIFICATIONS = [
    "Six Sigma Black Belt - Certified",
    "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
    "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
    "AWS IoT & SIMULINK - Amazon Web Services / MathWorks Certified",
]

PROJECT_SCAN_NOTES = [
    "# R05 Project Scan Notes",
    "",
    "## Source Reviewed",
    "",
    "Local archive: `/Users/kuldeepsingh/Downloads/HybridProjectpatent/NIDECG02+MICROVAST.zip`.",
    "",
    "The archive has damaged/incomplete central-directory metadata, but streaming extraction and local-header scans showed a real NIDECG02 + MICROVAST project folder with Simulink, MATLAB, DBC, MF4, generated target, calibration, and log artifacts.",
    "",
    "## Resume-Safe Evidence Extracted",
    "",
    "- Final Simulink model revisions: `G02_MICROVASTR03_FINAL.slx` and `G02_MICROVASTR10_FINAL.slx`.",
    "- Simulink metadata: R2018a model files with `HybridControlStrategy` harness entries, Stateflow XML, block-diagram XML, model workspace, and code dictionary artifacts.",
    "- MATLAB communication/control files: `VCVCCU_J1939.m`, `TorqueSpeedCtrl_LE_CAN.m`, `BMS_Charge_New.m`, `IXMV28.m`, `ZETTAJOULELTOBATTERY.m`, and `telematics.m`.",
    "- DBC evidence: `TATA1618_TSRTC.dbc` with 26 CAN messages and 211 signals across BMS, VCU command, motor status, telematics, EV faults, remote access, charge, precharge, voltage/current, torque/speed, and diagnostics.",
    "- Build/calibration evidence: generated `.mot` firmware outputs and `.a2l` measurement/calibration files under `Target_out`.",
    "- Test and calibration evidence: `.mf4` logs and code-log sheets showing assist/regen, temperature derate, cut-off voltage, and calibration update notes.",
    "",
    "## R05 Resume Integration",
    "",
    "- Added a page-1 project-verified proof section to reduce blank page space and make the resume more credible for automotive software lead screening.",
    "- Integrated only high-level, resume-safe facts; no source code, CAN IDs, proprietary logic, or copied implementation details are included in the resume.",
    "- Rebalanced R04's early forced page break into a fuller two-page layout.",
    "",
    "## Prior Benchmark Sources Retained",
    "",
    "- https://www.tesla.com/careers/search",
    "- https://dtici.daimlertruck.com/career/",
    "- https://www.tatatechnologies.com/us/solutions/",
    "- https://www.jaguarlandrovercareers.com/",
    "- https://www.avl.com/en/career/fields-interest/e-mobility",
    "- https://www.bosch.com/careers/",
    "- https://www.continental.com/en/career/",
    "",
]


def clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def set_font(run, name="Arial", size=9.0, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=2.2, line=1.02):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_rule(paragraph, color="D9D9D9", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.40)
    section.bottom_margin = Inches(0.40)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    for style_name, size in [("Normal", 8.35), ("List Bullet", 8.05)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(1.2)
        style.paragraph_format.line_spacing = 0.98


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=0, after=0.7, line=1.0)
    set_font(title.add_run(NAME.upper()), size=16.7, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=0.7, line=1.0)
    set_font(subtitle.add_run(HEADLINE), size=9.05, bold=True, color="1F4D78")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=2.4, line=1.0)
    set_font(contact.add_run(CONTACT), size=7.85, color="555555")


def add_continuation_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=2.5, line=1.0)
    set_font(p.add_run("KULDEEP SINGH | Project-Verified Automotive EV Software & Technical Lead | Page 2"), size=7.95, bold=True, color="555555")


def add_heading(doc, text, before=4.5):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=1.7, line=1.0)
    set_font(p.add_run(text), size=9.75, bold=True, color="1F4D78")
    add_bottom_rule(p)


def add_body(doc, text, size=8.28, after=1.4):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=0.99)
    set_font(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=8.12, after=0.55):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=after, line=0.98)
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_font(p.add_run(text), size=size)
    return p


def add_role_header(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0.6, after=0, line=1.0)
    set_font(p.add_run(ROLE), size=8.85, bold=True)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=0.9, line=1.0)
    set_font(p2.add_run(f"{COMPANY} | {DATES}"), size=7.95, color="555555")


def build_resume(out_dir: Path):
    doc = Document()
    configure_doc(doc)
    add_header(doc)

    add_heading(doc, "EXECUTIVE SUMMARY", before=1.4)
    add_body(doc, SUMMARY, size=8.28, after=1.1)

    add_heading(doc, "LEAD-ROLE FIT")
    for item in LEAD_ROLE_FIT:
        add_bullet(doc, item, size=8.15, after=0.35)

    add_heading(doc, "PROJECT-VERIFIED SIMULINK / STATEFLOW PROOF")
    for item in PROJECT_VERIFIED_PROOF:
        add_bullet(doc, item, size=8.02, after=0.3)

    add_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_header(doc)
    for item in EXPERIENCE:
        add_bullet(doc, item, size=8.0, after=0.25)

    add_heading(doc, "RELEASE & VALIDATION EVIDENCE")
    for item in RELEASE_VALIDATION_EVIDENCE:
        add_bullet(doc, item, size=7.98, after=0.18)

    add_heading(doc, "COMMERCIAL OUTCOMES")
    for item in COMMERCIAL_OUTCOMES:
        add_bullet(doc, item, size=7.98, after=0.18)

    doc.add_page_break()

    add_continuation_header(doc)
    add_heading(doc, "SELECTED PROGRAMMES & QUANTIFIED PROOF", before=0)
    for item in PROGRAMMES:
        add_bullet(doc, item, size=8.05, after=0.45)

    add_heading(doc, "BATTERY, CHARGING, EMBEDDED & VALIDATION PORTFOLIO")
    for item in PORTFOLIO:
        add_bullet(doc, item, size=8.03, after=0.42)

    add_heading(doc, "PROJECT IMPLEMENTATION DEPTH")
    for item in IMPLEMENTATION_DEPTH:
        add_bullet(doc, item, size=7.98, after=0.26)

    add_heading(doc, "TECHNICAL TOOLKIT & DELIVERY METHODS")
    for item in TECHNICAL_TOOLKIT:
        add_bullet(doc, item, size=8.0, after=0.35)

    add_heading(doc, "AUTOMOTIVE LEADERSHIP & DELIVERY")
    for item in LEADERSHIP_SIGNAL:
        add_bullet(doc, item, size=8.0, after=0.28)

    add_heading(doc, "APPLICATION SCOPE")
    for item in APPLICATION_SCOPE:
        add_bullet(doc, item, size=8.0, after=0.28)

    add_heading(doc, "CERTIFICATIONS")
    for item in CERTIFICATIONS:
        add_bullet(doc, item, size=8.05, after=0.25)

    add_heading(doc, "EDUCATION")
    add_body(doc, EDUCATION, size=8.08, after=0)

    docx_path = out_dir / f"{FILE_BASE}.docx"
    txt_path = out_dir / f"{FILE_BASE}.txt"
    doc.save(docx_path)

    txt_lines = [
        NAME.upper(),
        HEADLINE,
        CONTACT,
        "",
        "EXECUTIVE SUMMARY",
        SUMMARY,
        "",
        "LEAD-ROLE FIT",
        *[f"- {x}" for x in LEAD_ROLE_FIT],
        "",
        "PROJECT-VERIFIED SIMULINK / STATEFLOW PROOF",
        *[f"- {x}" for x in PROJECT_VERIFIED_PROOF],
        "",
        "PROFESSIONAL EXPERIENCE",
        ROLE,
        f"{COMPANY} | {DATES}",
        *[f"- {x}" for x in EXPERIENCE],
        "",
        "RELEASE & VALIDATION EVIDENCE",
        *[f"- {x}" for x in RELEASE_VALIDATION_EVIDENCE],
        "",
        "COMMERCIAL OUTCOMES",
        *[f"- {x}" for x in COMMERCIAL_OUTCOMES],
        "",
        "SELECTED PROGRAMMES & QUANTIFIED PROOF",
        *[f"- {x}" for x in PROGRAMMES],
        "",
        "BATTERY, CHARGING, EMBEDDED & VALIDATION PORTFOLIO",
        *[f"- {x}" for x in PORTFOLIO],
        "",
        "PROJECT IMPLEMENTATION DEPTH",
        *[f"- {x}" for x in IMPLEMENTATION_DEPTH],
        "",
        "TECHNICAL TOOLKIT & DELIVERY METHODS",
        *[f"- {x}" for x in TECHNICAL_TOOLKIT],
        "",
        "AUTOMOTIVE LEADERSHIP & DELIVERY",
        *[f"- {x}" for x in LEADERSHIP_SIGNAL],
        "",
        "APPLICATION SCOPE",
        *[f"- {x}" for x in APPLICATION_SCOPE],
        "",
        "CERTIFICATIONS",
        *[f"- {x}" for x in CERTIFICATIONS],
        "",
        "EDUCATION",
        EDUCATION,
        "",
    ]
    txt_path.write_text("\n".join(txt_lines), encoding="utf-8")
    return docx_path, txt_path


def build_readme(out_dir: Path):
    readme = [
        "# R05 Project-Verified Full Two-Page Resume",
        "",
        "Fuller automotive EV software and technical lead resume that integrates project-verified Simulink/Stateflow evidence from the NIDEC G02 + Microvast archive while staying resume-safe.",
        "",
        "## Files",
        f"- `{FILE_BASE}.docx`: editable Word resume",
        f"- `{FILE_BASE}.pdf`: rendered PDF resume",
        f"- `{FILE_BASE}.txt`: ATS/plain-text resume",
        "- `R05_Project_Scan_Notes.md`: high-level project scan notes",
        "",
        "## Why R05 Exists",
        "- R04 had visible blank space because of an early forced page break.",
        "- R05 fills both pages with useful, verified project proof instead of whitespace.",
        "- Simulink, Stateflow, CAN/J1939, DBC, MF4, `.mot`, `.a2l`, BMS charge, torque/speed, assist/regen, telematics, and diagnostics evidence is integrated naturally.",
        "- Proprietary implementation details and source code are not copied into the resume.",
        "",
        f"Version: `{VERSION}`",
        "",
    ]
    (out_dir / "README.md").write_text("\n".join(readme), encoding="utf-8")
    (out_dir / "R05_Project_Scan_Notes.md").write_text("\n".join(PROJECT_SCAN_NOTES), encoding="utf-8")


def main():
    for path in [OUTPUT_ROOT, REPO_R05, LOCAL_ROOT]:
        clean_dir(path)

    docx_path, txt_path = build_resume(OUTPUT_ROOT)
    build_readme(OUTPUT_ROOT)

    for file in OUTPUT_ROOT.iterdir():
        if file.is_file():
            shutil.copy2(file, REPO_R05 / file.name)
            shutil.copy2(file, LOCAL_ROOT / file.name)

    print(OUTPUT_ROOT)
    print(docx_path)
    print(txt_path)


if __name__ == "__main__":
    main()
