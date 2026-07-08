from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/kuldeepsingh/Documents/Codex/2026-07-08/en")
OUTPUT_ROOT = ROOT / "outputs" / "R07_HR_Reviewed_Technical_Manager_Resume"
REPO_ROOT = ROOT / "work" / "KuldeepSinghResumeFormat"
REPO_R07 = REPO_ROOT / "resumes" / "R07"
LOCAL_ROOT = Path("/Users/kuldeepsingh/Downloads/RESUME/R07_HR_Reviewed_Technical_Manager_Resume")

VERSION = "R07"
FILE_BASE = "Kuldeep_Singh_HR_Reviewed_EV_Technical_Manager_Resume_R07"

NAME = "Kuldeep Singh"
HEADLINE = "EV Technical Manager | Automotive Software Lead - eVCU/BMS, MBD, CAN, HV Battery"
CONTACT = "kuldeeppatel9125@gmail.com | +91 99448 76466 | Noida, India | linkedin.com/in/kuldeep9125"
COMPANY = "IX Energy Pvt. Ltd., Noida, India"
ROLE = "Technical Manager - Product Development"
DATES = "Jul 2018 - Present"
EDUCATION = "B.Tech - Mechanical Engineering, VIT University, Vellore, Tamil Nadu | 2014 - 2018 | First Class, 75%"

SUMMARY = (
    "EV Technical Manager and automotive software lead with 8 years delivering commercial EV, hybrid bus, HV battery, "
    "embedded electronics, and vehicle integration programmes from requirements through model-based software, supplier "
    "integration, validation evidence, and ICAT/ARAI homologation. Combines hands-on MATLAB/Simulink/Stateflow, "
    "eVCU/BMS/VCU logic, CAN/J1939 diagnostics, HV battery/charging, DCDC/OBC/PDU integration, and release discipline "
    "with cross-functional team and supplier leadership."
)

TECHNICAL_MANAGER_FIT = [
    "Lead-level ownership: requirements, architecture, SORs, supplier interfaces, model logic, calibration, integration, validation, certification evidence, and release readiness.",
    "Hands-on software depth: Simulink/Stateflow controls, eVCU/BMS/VCU functions, CAN/J1939/DBC communication, diagnostics, fault handling, safe-state logic, and calibration artefacts.",
    "Proven delivery: ARAI/ICAT-certified programmes, 16T P4 hybrid bus, 6.5T and 5T commercial EV platforms, 5 kWh to 300 kWh HV battery systems, and 25% fuel-efficiency validation.",
]

EXPERIENCE = [
    "Lead EV and P4 hybrid product development across ECU application software, HV battery engineering, power electronics, supplier development, vehicle integration, validation, and homologation readiness.",
    "Develop BMS/eVCU/VCU application software in MATLAB/Simulink/Stateflow for drive, charge, assist, regen, diagnostics, derating, fault response, precharge, active discharge, and safe-state behaviour.",
    "Define SORs, component specifications, CAN signal interfaces, test evidence, and supplier deliverables for motor, HV battery, PDU, DCDC, OBC, gearbox, controllers, HVAC, telematics, instrumentation, and embedded electronics.",
    "Architect certified commercial EV systems from requirement capture through ICAT/ARAI outputs, linking hardware, software, supplier readiness, vehicle validation, issue closure, and release documentation.",
    "Design and package 5 kWh to 300 kWh HV battery systems across 72V-800V LFP, NMC, LTO, and ultra-capacitor chemistries, including BMS/BMU logic, contactor/precharge, BTMS/HVAC, charging handshake, and pack integration.",
    "Use CAN diagnostics, DBC review, MF4/test-data analysis, bench and vehicle testing, calibration review, DFMEA/RCA, and MIL/SIL/HIL-ready model practices to reduce validation cycles and close issues.",
    "Manage internal engineering teams and external suppliers across mechanical, electrical, software, battery, power electronics, quality, and vehicle integration workstreams; report risk, readiness, and closure to leadership.",
]

MODEL_BASED_PROOF = [
    "Independently developed a project-verified hybrid control model in MATLAB/Simulink/Stateflow with final model revisions, generated target outputs, calibration artefacts, and bench/vehicle validation evidence.",
    "Implemented control logic for torque/speed requests, assist/regen, temperature derate, BMS state, SOC, charging limits, contactor/precharge, fault reset, MCU enable, active discharge, DTC/DM01-style diagnostics, and telematics status.",
    "Built CAN/J1939 and DBC-based communication across VCU, BMS, motor controller, charger, telematics, remote access, battery status, motor status, vehicle status, and EV fault signals.",
    "Used MF4 logs and code-log evidence to tune assist/regen calibration, temperature derate around 40-42 degC, cut-off voltage changes, charging limits, and validation-oriented release updates.",
]

PROGRAMMES = [
    "16T P4 hybrid bus: ICAT-certified platform with super-capacitor energy storage and validated 25% fuel-efficiency improvement; supported system integration, pilot fleet trials, validation evidence, and certification readiness.",
    "6.5T LCV/LPT commercial EV: ARAI-homologated platform with 80 kW powertrain, 53 kWh / 320V LFP battery, and 119 km range; contributed eVCU/BMS logic, DCDC/PDU/OBC integration, diagnostics, safety logic, and validation.",
    "5T commercial EV: supported 80 kW powertrain, 56 kWh / 310V NMC battery system, 140 km range target, HV battery integration, power electronics interfaces, supplier coordination, and vehicle validation.",
    "Hybrid controls programme: Simulink/Stateflow model with torque/speed, assist/regen, BMS charge, CAN/J1939, telematics, diagnostics, MF4 logs, MOT firmware output, and A2L calibration evidence.",
    "Multi-platform EV conversions: supported sedan and light commercial EV conversions across HV battery, charger, DCDC, controller, vehicle controls, diagnostics, CAN integration, and validation.",
]

PORTFOLIO = [
    "Battery and charging: 53 kWh/332V LFP truck pack, 11.5 kWh/332V LTO hybrid bus pack, 28 kWh NMC pack, 0.5 kWh/400V ultra-capacitor pack, AC/DC charging controller logic, charger/BMS/VCU coordination, HV-LV handshake, contactor/precharge, and fault response.",
    "Control and diagnostics: torque request, speed limit, gear, brake, accelerator, MCU enable, fault reset, active discharge, FailGrade, motor speed/torque/current, DC voltage/current, DTC/SPN-style diagnostics, remote access, and vehicle status interfaces.",
    "Embedded electronics: LV/HV PDU, STM32F4/F1 instrument cluster, Quectel dual-CAN telematics unit, ultra-capacitor cell monitoring/control system, battery-related software interfaces, and CAN-based diagnostics.",
    "Validation and homologation: ARAI/ICAT evidence, DBC/CAN review, MF4 logs, code-log updates, assist/regen calibration, temperature derate tuning, bench/vehicle testing, validation reports, and release documentation.",
]

TECHNICAL_TOOLKIT = [
    "MBD/software: MATLAB, Simulink, Stateflow, Model-Based Design, Embedded C/C++, generated build outputs, A2L, MOT, calibration/measurement artefacts.",
    "Networks/diagnostics: CAN, J1939, DBC, ECOCAN-style scripts, Vector CANoe, Peak CAN, Bus Master, CSS Electronics, MF4 logs, UDS-aware diagnostics, DTC/SPN concepts.",
    "EV systems: eVCU, VCU, BMS, HV battery, LFP, NMC, LTO, ultra-capacitor, DCDC, OBC, PDU, charger, contactor, precharge, BTMS/HVAC, motor controller, e-drive/EDU, hybrid powertrain.",
]

QUALITY_LEADERSHIP = [
    "Run programme cadence from vehicle objectives to software requirements, interface freeze, supplier DVP, vehicle validation, release notes, and issue closure.",
    "Requirements and release governance: maintain SORs, component specifications, validation evidence, supplier acceptance criteria, and release documentation from requirement to vehicle sign-off.",
    "Translate test data into model changes, CAN/DBC updates, calibration decisions, supplier actions, risk closure, and certification evidence.",
    "Safety and quality methods: apply ISO 26262 awareness, IATF 16949 discipline, DFMEA/RCA, Six Sigma, DMAIC, SPC, control charts, and production-readiness review habits.",
    "Supplier leadership: align motor controller, HV battery/BMS, charger, DCDC/OBC/PDU, telematics, instrumentation, and embedded electronics suppliers against test evidence and issue closure.",
    "Team leadership: guide mechanical, electrical, embedded software, battery, power electronics, quality, and vehicle integration teams with clear risk, readiness, and closure reporting.",
    "Decision balance: manage range, performance, thermal limits, diagnostics coverage, safety response, supplier readiness, cost, timing, serviceability, and manufacturing constraints.",
    "Forward-looking readiness: AUTOSAR basics, SDV awareness, MIL/SIL/HIL readiness, model/code review mindset, and validation-data-driven release decisions.",
]

CERTIFICATIONS = [
    "Six Sigma Black Belt - Certified",
    "ISO 26262 Functional Safety for Road Vehicles - TUV SUD Certified",
    "IATF 16949:2016 Automotive Quality Management - TUV SUD Certified",
    "AWS IoT & SIMULINK - Amazon Web Services / MathWorks Certified",
]

HR_REVIEW_NOTES = [
    "# R07 HR Review Notes",
    "",
    "## HR Review Finding",
    "",
    "R06 was technically strong but had too much duplication for a Technical Manager resume. The same model, CAN, validation, supplier, and release proof appeared across multiple sections, making the resume feel dense and slightly repetitive.",
    "",
    "## Duplicates Reduced",
    "",
    "- Removed separate `Application Scope`; target-role language is now in the headline and manager-fit bullets.",
    "- Removed separate `Project Implementation Depth`; the strongest evidence is now folded into one model-based proof section.",
    "- Merged `Release & Validation Evidence` and `Commercial Outcomes` into current-role and programme proof.",
    "- Kept project identifiers privacy-clean; no supplier/project codes or file names are used in the resume.",
    "",
    "## R07 Structure",
    "",
    "- Page 1: summary, manager fit, current role, model-based controls proof, and strongest delivery proof.",
    "- Page 2: additional platform proof, technical portfolio, toolkit/process, certifications, education.",
    "- Visual intent: fewer section headers, stronger skim path, less repeated wording, and a cleaner technical-manager story.",
    "",
]


def clean_dir(path: Path) -> None:
    if path.exists():
        shutil.rmtree(path)
    path.mkdir(parents=True, exist_ok=True)


def set_font(run, name="Arial", size=9.0, bold=False, italic=False, color="000000"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_spacing(paragraph, before=0, after=2.4, line=1.04):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_rule(paragraph, color="D9D9D9", size="6", space="1"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.40)
    section.bottom_margin = Inches(0.40)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)

    for style_name, size in [("Normal", 8.9), ("List Bullet", 8.65)]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Arial")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial")
        style.font.size = Pt(size)
        style.paragraph_format.space_after = Pt(1.9)
        style.paragraph_format.line_spacing = 1.04


def add_header(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(title, before=0, after=0.7, line=1.0)
    set_font(title.add_run(NAME.upper()), size=17.1, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(subtitle, after=0.7, line=1.0)
    set_font(subtitle.add_run(HEADLINE), size=9.25, bold=True, color="1F4D78")

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(contact, after=2.4, line=1.0)
    set_font(contact.add_run(CONTACT), size=8.0, color="555555")


def add_continuation_header(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, before=0, after=3.4, line=1.0)
    set_font(p.add_run("KULDEEP SINGH | EV Technical Manager | Page 2"), size=7.95, bold=True, color="555555")


def add_heading(doc, text, before=5.4):
    p = doc.add_paragraph()
    set_spacing(p, before=before, after=2.25, line=1.0)
    set_font(p.add_run(text), size=10.35, bold=True, color="1F4D78")
    add_bottom_rule(p)


def add_body(doc, text, size=8.48, after=1.4):
    p = doc.add_paragraph()
    set_spacing(p, after=after, line=1.04)
    set_font(p.add_run(text), size=size)
    return p


def add_bullet(doc, text, size=8.3, after=0.55):
    p = doc.add_paragraph(style="List Bullet")
    set_spacing(p, after=after, line=1.04)
    p.paragraph_format.left_indent = Inches(0.20)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    set_font(p.add_run(text), size=size)
    return p


def add_role_header(doc):
    p = doc.add_paragraph()
    set_spacing(p, before=0.6, after=0, line=1.0)
    set_font(p.add_run(ROLE), size=9.05, bold=True)
    p2 = doc.add_paragraph()
    set_spacing(p2, before=0, after=0.9, line=1.0)
    set_font(p2.add_run(f"{COMPANY} | {DATES}"), size=8.12, color="555555")


def build_resume(out_dir: Path):
    doc = Document()
    configure_doc(doc)
    add_header(doc)

    add_heading(doc, "EXECUTIVE SUMMARY", before=1.4)
    add_body(doc, SUMMARY, size=8.7, after=1.4)

    add_heading(doc, "TECHNICAL MANAGER FIT")
    for item in TECHNICAL_MANAGER_FIT:
        add_bullet(doc, item, size=8.55, after=0.55)

    add_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_header(doc)
    for item in EXPERIENCE:
        add_bullet(doc, item, size=8.42, after=0.45)

    add_heading(doc, "MODEL-BASED CONTROLS & VALIDATION PROOF")
    for item in MODEL_BASED_PROOF:
        add_bullet(doc, item, size=8.45, after=0.5)

    add_heading(doc, "SELECTED DELIVERY PROOF")
    for item in PROGRAMMES[:3]:
        add_bullet(doc, item, size=8.48, after=0.58)

    doc.add_page_break()

    add_continuation_header(doc)
    add_heading(doc, "SELECTED PROGRAMMES & QUANTIFIED PROOF", before=0)
    for item in PROGRAMMES[3:]:
        add_bullet(doc, item, size=8.48, after=0.58)

    add_heading(doc, "TECHNICAL PORTFOLIO")
    for item in PORTFOLIO:
        add_bullet(doc, item, size=8.45, after=0.55)

    add_heading(doc, "TOOLKIT, PROCESS & LEADERSHIP")
    for item in TECHNICAL_TOOLKIT:
        add_bullet(doc, item, size=8.42, after=0.5)

    add_heading(doc, "QUALITY, SAFETY & SUPPLIER LEADERSHIP")
    for item in QUALITY_LEADERSHIP:
        add_bullet(doc, item, size=8.42, after=0.48)

    add_heading(doc, "CERTIFICATIONS")
    for item in CERTIFICATIONS:
        add_bullet(doc, item, size=8.45, after=0.38)

    add_heading(doc, "EDUCATION")
    add_body(doc, EDUCATION, size=8.45, after=0)

    docx_path = out_dir / f"{FILE_BASE}.docx"
    txt_path = out_dir / f"{FILE_BASE}.txt"
    doc.save(docx_path)

    txt_lines = [
        NAME.upper(),
        HEADLINE,
        CONTACT,
        "",
        "EXECUTIVE SUMMARY",
        SUMMARY,
        "",
        "TECHNICAL MANAGER FIT",
        *[f"- {x}" for x in TECHNICAL_MANAGER_FIT],
        "",
        "PROFESSIONAL EXPERIENCE",
        ROLE,
        f"{COMPANY} | {DATES}",
        *[f"- {x}" for x in EXPERIENCE],
        "",
        "MODEL-BASED CONTROLS & VALIDATION PROOF",
        *[f"- {x}" for x in MODEL_BASED_PROOF],
        "",
        "SELECTED DELIVERY PROOF",
        *[f"- {x}" for x in PROGRAMMES[:3]],
        "",
        "SELECTED PROGRAMMES & QUANTIFIED PROOF",
        *[f"- {x}" for x in PROGRAMMES[3:]],
        "",
        "TECHNICAL PORTFOLIO",
        *[f"- {x}" for x in PORTFOLIO],
        "",
        "TOOLKIT, PROCESS & LEADERSHIP",
        *[f"- {x}" for x in TECHNICAL_TOOLKIT],
        "",
        "QUALITY, SAFETY & SUPPLIER LEADERSHIP",
        *[f"- {x}" for x in QUALITY_LEADERSHIP],
        "",
        "CERTIFICATIONS",
        *[f"- {x}" for x in CERTIFICATIONS],
        "",
        "EDUCATION",
        EDUCATION,
        "",
    ]
    txt_path.write_text("\n".join(txt_lines), encoding="utf-8")
    return docx_path, txt_path


def build_readme(out_dir: Path):
    readme = [
        "# R07 HR-Reviewed Technical Manager Resume",
        "",
        "HR-reviewed, duplicate-reduced two-page resume for EV Technical Manager and automotive software lead roles.",
        "",
        "## Files",
        f"- `{FILE_BASE}.docx`: editable Word resume",
        f"- `{FILE_BASE}.pdf`: rendered PDF resume",
        f"- `{FILE_BASE}.txt`: ATS/plain-text resume",
        "- `R07_HR_Review_Notes.md`: duplicate-reduction and structure notes",
        "",
        "## Why R07 Exists",
        "- R06 was strong but repeated the same proof across too many sections.",
        "- R07 removes duplicate sections, keeps project identifiers privacy-clean, and improves HR/technical-manager readability.",
        "- The latest resume keeps Simulink, Stateflow, CAN/J1939, DBC, MF4, MOT, A2L, BMS charge, torque/speed, assist/regen, telematics, diagnostics, certification, and leadership evidence without repeating it.",
        "",
        f"Version: `{VERSION}`",
        "",
    ]
    (out_dir / "README.md").write_text("\n".join(readme), encoding="utf-8")
    (out_dir / "R07_HR_Review_Notes.md").write_text("\n".join(HR_REVIEW_NOTES), encoding="utf-8")


def main():
    for path in [OUTPUT_ROOT, REPO_R07, LOCAL_ROOT]:
        clean_dir(path)

    docx_path, txt_path = build_resume(OUTPUT_ROOT)
    build_readme(OUTPUT_ROOT)

    for file in OUTPUT_ROOT.iterdir():
        if file.is_file():
            shutil.copy2(file, REPO_R07 / file.name)
            shutil.copy2(file, LOCAL_ROOT / file.name)

    print(OUTPUT_ROOT)
    print(docx_path)
    print(txt_path)


if __name__ == "__main__":
    main()
